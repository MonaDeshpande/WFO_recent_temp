import pandas as pd
from sqlalchemy import create_engine
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import sqlalchemy
import re

# Database connection parameters (update with your actual details)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag and calculated KPI
TAG_UNITS = {
    'FT-01': 'kg/h', 'FT-02': 'kg/h', 'FT-03': 'kg/h', 'FT-05': 'kg/h', 'FT-06': 'kg/h',
    'FT-09': 'kg/h', 'FT-61': 'kg/h', 'FT-62': 'kg/h', 'FI-103': 'kg/h', 'FI-202': 'kg/h',
    'TI-11': 'degC', 'TI-13': 'degC', 'TI-14': 'degC', 'TI-15': 'degC', 'TI-16': 'degC',
    'TI-17': 'degC', 'TI-18': 'degC', 'TI-19': 'degC', 'TI-20': 'degC', 'TI-21': 'degC',
    'TI-22': 'degC', 'TI-23': 'degC', 'TI-24': 'degC', 'TI-25': 'degC', 'TI-26': 'degC',
    'TI-27': 'degC', 'TI-28': 'degC', 'TI-29': 'degC', 'TI-30': 'degC', 'TI-72A': 'degC',
    'TI-72B': 'degC', 'PTT-02': 'mmHg', 'PTB-02': 'mmHg',
    # Calculated KPIs (using descriptive keys for easy lookup in report)
    'DIFFERENTIAL_PRESSURE': 'mmHg',
    'LI-03': '%',
    'REBOILER_HEAT_DUTY': 'kW',  # Heat duty now fetched, IQR analysis applied
    'CONDENSER_HEAT_DUTY': 'kW', # Heat duty now fetched, IQR analysis applied
    'REFLUX_RATIO': '',  # Dimensionless
    'MATERIAL_BALANCE_ERROR': '%',
    'NAPHTHALENE_LOSS_PERCENTAGE': '%'
}

# File paths for saving generated plots and report
output_report_path = "C-02_Analysis_Report.docx"
output_temp_plot_path = "C-02_Temperature_Profile.png"
output_dp_plot_path = "C-02_Differential_Pressure.png"
output_trends_plot_path = "C-02_Daily_Trends.png"
output_naphthalene_vs_reflux_plot_path = "C-02_Naphthalene_vs_Reflux.png"
output_naphthalene_vs_reboiler_temp_plot_path = "C-02_Naphthalene_vs_Reboiler_Temp.png"
output_feed_vs_dp_plot_path = "C-02_Feed_vs_DP.png"

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine):
    """Retrieves specific SCADA data for the C-02 column and related streams."""
    try:
        # Define all desired columns, including the new pre-calculated heat duty columns
        desired_columns = [
            "DateAndTime", "FT-01", "FT-02", "FT-03", "FT-05", "FT-06", "FT-09", "FT-61", "FT-62", "TI-11", "TI-13", "TI-14", "TI-15", "TI-16",
            "TI-17", "TI-18", "TI-19", "TI-20", "TI-21", "TI-22", "TI-23", "TI-24", "TI-25", "TI-26",
            "TI-27", "TI-28", "TI-29", "TI-30", "TI-72A", "TI-72B", "PTT-02", "PTB-02", "LI-03",
            "FI-103", "FI-202",
            "HeatDuty_C-02_Reboiler",  # Fetched Reboiler duty
            "HeatDuty_C-02_Condenser"  # Fetched Condenser duty
        ]
        
        start_date = '2025-09-03 00:00:00'
        end_date = '2025-09-30 00:00:00'

        inspector = sqlalchemy.inspect(engine)
        columns = inspector.get_columns('data_cleaning_with_report')
        column_names = [col['name'] for col in columns]
        
        final_columns = []
        
        # Mapping for special columns that must be renamed after fetching
        rename_map = {}
        
        # Match desired columns (case and hyphen insensitive) with database columns
        for d_col in desired_columns:
            # Normalize desired column name for comparison (e.g., 'FT-01' -> 'ft01')
            d_col_normalized = d_col.lower().replace('-', '').replace('_', '')
            for db_col in column_names:
                # Normalize database column name for comparison (e.g., 'ft_01' -> 'ft01')
                db_col_normalized = db_col.lower().replace('-', '').replace('_', '')
                if d_col_normalized == db_col_normalized:
                    final_columns.append(f'"{db_col}"')
                    
                    # Prepare for renaming the heat duty columns to the standardized internal names
                    if 'heatduty' in d_col_normalized:
                            if 'reboiler' in d_col_normalized:
                                rename_map[db_col.upper().replace('-', '_')] = 'REBOILER_HEAT_DUTY'
                            elif 'condenser' in d_col_normalized:
                                rename_map[db_col.upper().replace('-', '_')] = 'CONDENSER_HEAT_DUTY'

                    break
            
        if not final_columns:
            print("Error: No matching columns found. Data retrieval failed.")
            return None, start_date, end_date

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM data_cleaning_with_report
        WHERE "DateAndTime" BETWEEN '{start_date}' AND '{end_date}'
        ORDER BY "DateAndTime";
        """
        
        df = pd.read_sql(query, engine)
        
        # Standardize column names in the DataFrame to uppercase and use underscores
        df.columns = [col.upper().replace('-', '_') for col in df.columns]
        
        # Explicitly rename the fetched Heat Duty columns to the standardized KPI names
        df.rename(columns=rename_map, inplace=True)
        
        df['DATEANDTIME'] = pd.to_datetime(df['DATEANDTIME'])
        print("SCADA data for process streams retrieved successfully.")
        return df, start_date, end_date
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None, None, None

def get_composition_data():
    """
    Simulates reading composition data from a lab analysis report.
    Returns a dictionary of compositions for each component at specific sample points.
    """
    try:
        composition_data = {
            'Naphthalene': {
                'P-01': 0.1,    # C-00 Feed
                'C-01-B': 0.02, # C-01 Bottoms (Light ends)
                'C-02-T': 0.08, # C-02 Tops (Product)
            },
            'Moisture': {
                'P-01': 0.15 # Moisture content in the C-00 Feed (P-01)
            }
        }
        return composition_data
    except Exception as e:
        print(f"Error simulating composition data: {e}. Using default values.")
        return None

def clean_outliers(series, tag, threshold=3):
    """
    Applies 3-sigma capping for general data tags.
    Caps values in a series that are outside the mean +/- (threshold * std_dev).
    """
    data = series.dropna()
    if data.empty:
        return series, 0
    
    mean = data.mean()
    std = data.std()
    lower_bound = mean - threshold * std
    upper_bound = mean + threshold * std
    
    outlier_count = series[(series < lower_bound) | (series > upper_bound)].count()
    
    # Apply capping
    series_cleaned = np.clip(series, lower_bound, upper_bound)
    return series_cleaned, outlier_count

def iqr_filter_and_get_cleaned_data(series, column_name):
    """
    Applies the IQR method (1.5 * IQR) to detect and filter outliers
    from a specified data series. Returns the cleaned series and a report dict.
    The 'Cleaned_Average' is the Median (50th percentile) of the filtered data.
    """
    series_nn = series.dropna()
    if series_nn.empty:
        return series, {'Column': column_name, 'Total_Points': 0, 'Outlier_Count': 0, 'Percentage_Removed': 0.0, 'Median_Value': np.nan, 'Q1': np.nan, 'Q3': np.nan}

    # 1. Calculate Q1 (25th percentile) and Q3 (75th percentile)
    Q1 = series_nn.quantile(0.25)
    Q3 = series_nn.quantile(0.75)
    
    # 2. Calculate IQR and Bounds
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    
    # 3. Filter the Data (Outliers are replaced with NaN)
    # Note: We use the bounds on the original series, but the median/quartiles 
    # for reporting are based on the original non-NaN series for accuracy.
    df_cleaned = series.apply(lambda x: x if (x >= lower_bound and x <= upper_bound) else np.nan)
    
    # 4. Calculate metrics for the report
    original_count = len(series_nn) # Total number of non-NaN points analyzed
    cleaned_count = df_cleaned.dropna().count()
    outlier_count = original_count - cleaned_count
    
    # Calculate the percentage of points removed
    percentage_removed = (outlier_count / original_count) * 100 if original_count > 0 else 0

    report = {
        'Column': column_name,
        'Total_Points': original_count, 
        'Outlier_Count': outlier_count,
        'Percentage_Removed': percentage_removed, 
        'Median_Value': series_nn.quantile(0.50), # The robust average for the report
        'Q1': Q1,
        'Q3': Q3
    }
    return df_cleaned, report

def perform_analysis(df):
    """
    Performs key calculations for the C-02 column and the overall process,
    including staged material balances, KPI calculation, and outlier cleaning.
    """
    if df is None or df.empty:
        return {}, df, {}, {}

    analysis_results = {}
    composition_data = get_composition_data()
    outliers_removed_3sigma = {}
    iqr_outlier_summary = {}
    
    # --- 1. Initial Data Cleaning & Calculation of Intermediate KPIs (DP and RR) ---
    
    # Convert data to numeric (coercing errors to NaN)
    tags_to_numeric = [col for col in df.columns if col not in ['DATEANDTIME', 'DATE']]
    for col in tags_to_numeric:
        df.loc[:, col] = pd.to_numeric(df[col], errors='coerce')

    # Differential Pressure (DP) Calculation - Required before cleaning DP
    if 'PTT_02' in df.columns and 'PTB_02' in df.columns:
        df['DIFFERENTIAL_PRESSURE'] = df['PTB_02'] - df['PTT_02']
        
    # Reflux Ratio (C-02) Calculation - Required before cleaning RR
    if 'FT_09' in df.columns and 'FT_03' in df.columns:
        df['REFLUX_RATIO'] = df['FT_09'] / df['FT_03']
        df.loc[df['FT_03'] == 0, 'REFLUX_RATIO'] = np.nan
        df['REFLUX_RATIO'] = df['REFLUX_RATIO'].abs()
        
    # --- 2. Outlier Removal ---

    # Tags for standard 3-sigma capping (flows, temperatures, DP)
    tags_for_3_sigma = ['FT_01', 'FT_02', 'FT_03', 'FT_06', 'FT_09', 'DIFFERENTIAL_PRESSURE', 'TI_72B', 'TI_13'] # Added TI-13 for temp check
    
    for tag in tags_for_3_sigma:
        if tag in df.columns:
            df[tag], count = clean_outliers(df[tag], tag)
            if count > 0:
                outliers_removed_3sigma[tag] = count

    # Tags for specific IQR filtering (Heat Duties and Reflux Ratio)
    # Reflux Ratio is added here as requested to use the robust median average.
    tags_for_iqr = ['REBOILER_HEAT_DUTY', 'CONDENSER_HEAT_DUTY', 'REFLUX_RATIO']
    
    for tag in tags_for_iqr:
        if tag in df.columns:
            # Apply IQR filtering
            cleaned_series, report = iqr_filter_and_get_cleaned_data(df[tag], tag)
            
            # Update the DataFrame with the cleaned series (outliers replaced by NaN)
            df[tag] = cleaned_series
            
            # Store the IQR results for the report
            iqr_outlier_summary[tag] = report
            
            # Update the main analysis_results with the CLEANED IQR average (Median)
            analysis_results[f'Average {tag} (IQR Cleaned)'] = report['Median_Value']
            # Store Q1 and Q3 for dedicated reporting in the docx file
            analysis_results[f'{tag} Q1'] = report['Q1']
            analysis_results[f'{tag} Q3'] = report['Q3']

    # --- 3. Final KPI Calculation (using cleaned data) ---
    
    # Calculate average flows for the period (using mean, which ignores NaNs)
    avg_flows = {tag.replace('-', '_').upper(): df[tag.replace('-', '_').upper()].mean()  
                 for tag in ['FT-01', 'FT-61', 'FT-62', 'FT-05', 'FT-02', 'FT-03', 'FT-06', 'FT-09']}

    # --- Staged Material Balance Calculations (Same logic, now using cleaned flows) ---
    input_flow_c00 = avg_flows.get('FT_01', 0)
    output_flow_c00 = avg_flows.get('FT_62', 0) + avg_flows.get('FT_61', 0)
    c00_balance_error = ((input_flow_c00 - output_flow_c00) / input_flow_c00) * 100 if input_flow_c00 > 0 else 0
    analysis_results['C-00 Overall Material Balance Error (%)'] = abs(c00_balance_error)

    # C-01 Naphthalene Loss Calculation (Uses assumed/simulated composition data)
    if composition_data and 'Naphthalene' in composition_data and 'Moisture' in composition_data:
        moisture_removed = composition_data['Moisture'].get('P-01', 0)
        c01_feed_comp_naphthalene = composition_data['Naphthalene'].get('P-01', 0) / (1 - moisture_removed)
        
        naphthalene_mass_in_c01_feed = avg_flows.get('FT_62', 0) * c01_feed_comp_naphthalene
        naphthalene_mass_in_c01_bottom = avg_flows.get('FT_05', 0) * composition_data['Naphthalene'].get('C-01-B', 0)
        
        if naphthalene_mass_in_c01_feed > 0:
            naphthalene_loss_percent_c01 = (naphthalene_mass_in_c01_bottom / naphthalene_mass_in_c01_feed) * 100
        else:
            naphthalene_loss_percent_c01 = 0
            
        analysis_results['Naphthalene Loss in C-01 (%)'] = naphthalene_loss_percent_c01
        
        if naphthalene_loss_percent_c01 > 2:
            analysis_results['C-01 Naphthalene Loss Status'] = "ALERT: Naphthalene loss is above the 2% limit."
        else:
            analysis_results['C-01 Naphthalene Loss Status'] = "Naphthalene loss is within acceptable limits."

    # C-02 Material Balance 
    if all(tag in df.columns for tag in ['FT_02', 'FT_03', 'FT_06']):
        feed_flow_avg = avg_flows.get('FT_02', 0)
        top_product_flow_avg = avg_flows.get('FT_03', 0)
        bottom_product_flow_avg = avg_flows.get('FT_06', 0)
        
        if feed_flow_avg > 0:
            material_balance_error = ((feed_flow_avg - (top_product_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
            analysis_results['C-02 Overall Material Balance Error (%)'] = abs(material_balance_error)

    # Differential Pressure Averages
    if 'DIFFERENTIAL_PRESSURE' in df.columns:
        analysis_results['Average Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DIFFERENTIAL_PRESSURE'].max()
        
    # Energy Balance Averages (Using original descriptive names and the IQR cleaned values)
    analysis_results['Average Reboiler Heat Duty'] = analysis_results.pop('Average REBOILER_HEAT_DUTY (IQR Cleaned)', 'N/A')
    analysis_results['Average Condenser Heat Duty'] = analysis_results.pop('Average CONDENSER_HEAT_DUTY (IQR Cleaned)', 'N/A')
    analysis_results['Average Reflux Ratio (IQR Median)'] = analysis_results.pop('Average REFLUX_RATIO (IQR Cleaned)', 'N/A')
    
    analysis_results['Reboiler Temp (TI-72B)'] = df['TI_72B'].mean() # Using cleaned TI_72B

    # Create new columns for plotting based on impurity in top product
    if 'FT_03' in df.columns and composition_data and 'Naphthalene' in composition_data and 'C-02-T' in composition_data['Naphthalene']:
        naphthalene_comp_in_top = composition_data['Naphthalene']['C-02-T']
        df['NAPHTHALENE_IN_C02_TOP_PROD_MASS'] = df['FT_03'] * naphthalene_comp_in_top
        analysis_results['Naphthalene in C-02 Top Product (%)'] = naphthalene_comp_in_top * 100
        
    return analysis_results, df, outliers_removed_3sigma, iqr_outlier_summary

def generate_plots(df):
    """Generates and saves temperature profile, DP, and energy plots."""
    
    if df is None or df.empty:
        print("Dataframe is empty, cannot generate plots.")
        return

    # Helper function to generate a plot
    def plot_and_save(x_data, y_data, title, xlabel, ylabel, filename, is_scatter=False):
        plt.figure(figsize=(10, 6))
        # Drop rows with NaN in data used for plotting
        plot_df = pd.DataFrame({'x': x_data, 'y': y_data}).dropna()
        
        if plot_df.empty:
            print(f"Skipping plot {filename}: Data is empty after NaN removal.")
            plt.close()
            return

        if is_scatter:
            plt.scatter(plot_df['x'], plot_df['y'], alpha=0.5)
        else:
            plt.plot(plot_df['x'], plot_df['y'], alpha=0.7)
            
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(filename)
        plt.close()
        print(f"Plot saved to {filename}")
        
    # Temperature Profile Plot
    if 'DATEANDTIME' in df.columns:
        # Select temperature tags across the column
        temp_tags = ['TI_13', 'TI_14', 'TI_15', 'TI_16', 'TI_17', 'TI_18', 'TI_19', 'TI_20', 'TI_21', 'TI_22', 'TI_23', 'TI_24', 'TI_25']
        
        plt.figure(figsize=(10, 6))
        df.sort_values(by='DATEANDTIME', inplace=True)
        x_axis = df['DATEANDTIME']
        
        for tag_underscore in temp_tags:
            if tag_underscore in df.columns:
                tag_hyphen = tag_underscore.replace('_', '-')
                plt.plot(x_axis, df[tag_underscore], label=tag_hyphen, alpha=0.7)
        
        plt.title("C-02 Column Temperature Profile Over Time")
        plt.xlabel("Date and Time")
        plt.ylabel(f"Temperature ({TAG_UNITS['TI-13']})")
        plt.legend(ncol=3, loc='lower center', bbox_to_anchor=(0.5, -0.25))
        plt.grid(True)
        plt.tight_layout(rect=[0, 0.1, 1, 1]) # Adjust layout for legend
        plt.savefig(output_temp_plot_path)
        plt.close()
        print(f"Temperature profile plot saved to {output_temp_plot_path}")

    # Differential Pressure Plot
    if 'DIFFERENTIAL_PRESSURE' in df.columns:
        plot_and_save(df['DATEANDTIME'], df['DIFFERENTIAL_PRESSURE'], 
                      "C-02 Differential Pressure Over Time", "Date and Time", 
                      f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})", 
                      output_dp_plot_path)

    # Daily Trends Plot
    if all(col in df.columns for col in ['DATEANDTIME', 'FT_03', 'TI_28', 'DIFFERENTIAL_PRESSURE']):
        df['DATE'] = df['DATEANDTIME'].dt.date
        # Ensure that mean aggregation works on cleaned data (which may have NaNs)
        daily_trends = df.groupby('DATE').agg({
            'FT_03': 'mean',
            'TI_28': 'mean', 
            'DIFFERENTIAL_PRESSURE': 'mean'
        }).reset_index()
        
        plt.figure(figsize=(12, 8))
        plt.plot(daily_trends['DATE'], daily_trends['FT_03'], label=f"Avg Top Product Flow ({TAG_UNITS['FT-03']})")
        plt.plot(daily_trends['DATE'], daily_trends['TI_28'], label=f"Avg Top Product Temp ({TAG_UNITS['TI-28']})")
        plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})")
        
        plt.title("C-02 Daily Trends")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.legend()
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(output_trends_plot_path)
        plt.close()
        print(f"Daily trends plot saved to {output_trends_plot_path}")

    # Naphthalene Loss vs. Reflux Ratio (C-02)
    if 'NAPHTHALENE_IN_C02_TOP_PROD_MASS' in df.columns and 'REFLUX_RATIO' in df.columns:
        plot_and_save(df['REFLUX_RATIO'], df['NAPHTHALENE_IN_C02_TOP_PROD_MASS'], 
                      "C-02 Naphthalene Loss vs. Reflux Ratio", 
                      "Reflux Ratio (R=FT-09/FT-03)", "Naphthalene Mass in Top Product (kg/h)", 
                      output_naphthalene_vs_reflux_plot_path, is_scatter=True)

    # Naphthalene Loss vs. Reboiler Temperature (C-02)
    if 'NAPHTHALENE_IN_C02_TOP_PROD_MASS' in df.columns and 'TI_72B' in df.columns:
        plot_and_save(df['TI_72B'], df['NAPHTHALENE_IN_C02_TOP_PROD_MASS'], 
                      "C-02 Naphthalene Loss vs. Reboiler Temperature (TI-72B)", 
                      f"Reboiler Temperature ({TAG_UNITS['TI-72B']})", "Naphthalene Mass in Top Product (kg/h)", 
                      output_naphthalene_vs_reboiler_temp_plot_path, is_scatter=True)
                        
    # Feed vs. Differential Pressure (C-02)
    if 'FT_02' in df.columns and 'DIFFERENTIAL_PRESSURE' in df.columns:
        plot_and_save(df['FT_02'], df['DIFFERENTIAL_PRESSURE'], 
                      "C-02 Feed Flow (FT-02) vs. Differential Pressure (DP)", 
                      f"Feed Flow (FT-02) ({TAG_UNITS['FT-02']})", 
                      f"Differential Pressure ({TAG_UNITS['DIFFERENTIAL_PRESSURE']})", 
                      output_feed_vs_dp_plot_path, is_scatter=True)
        
def generate_word_report(analysis_results, df, outliers_removed_3sigma, iqr_outlier_summary, start_date, end_date):
    """Creates a detailed analysis report in a Word document, including detailed outlier data."""
    doc = Document()
    doc.add_heading('C-02 Light Oil Recovery Column Analysis Report', 0)
    doc.add_paragraph(f"Analysis Period: {start_date} to {end_date}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Map descriptive result keys to TAG_UNITS keys for unit lookup
    KPI_UNIT_MAP = {
        'Average Reflux Ratio (IQR Median)': 'REFLUX_RATIO',
        'C-02 Overall Material Balance Error (%)': 'MATERIAL_BALANCE_ERROR',
        'Average Differential Pressure': 'DIFFERENTIAL_PRESSURE',
        'Maximum Differential Pressure': 'DIFFERENTIAL_PRESSURE',
        'Average Reboiler Heat Duty': 'REBOILER_HEAT_DUTY',
        'Average Condenser Heat Duty': 'CONDENSER_HEAT_DUTY',
        'Naphthalene in C-02 Top Product (%)': 'NAPHTHALENE_LOSS_PERCENTAGE',
        'Naphthalene Loss in C-01 (%)': 'NAPHTHALENE_LOSS_PERCENTAGE',
        'C-00 Overall Material Balance Error (%)': 'MATERIAL_BALANCE_ERROR',
        'Reboiler Temp (TI-72B)': 'TI-72B'
    }

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_text = ""
    
    # Use the IQR-cleaned averages for the summary
    avg_reboiler_duty = analysis_results.get('Average Reboiler Heat Duty')
    avg_condenser_duty = analysis_results.get('Average Condenser Heat Duty')

    if isinstance(avg_reboiler_duty, (float, int)):
        summary_text += f"The **Reboiler Heat Duty (IQR Cleaned)** averaged {avg_reboiler_duty:.2f} kW. "
    if isinstance(avg_condenser_duty, (float, int)):
        summary_text += f"The **Condenser Heat Duty (IQR Cleaned)** averaged {avg_condenser_duty:.2f} kW. "
    
    c02_error = analysis_results.get('C-02 Overall Material Balance Error (%)')
    if isinstance(c02_error, (float, int)):
        summary_text += f"The C-02 column had a material balance error of {c02_error:.2f}%. "
    
    avg_reflux = analysis_results.get('Average Reflux Ratio (IQR Median)')
    reflux_q1 = analysis_results.get('REFLUX_RATIO Q1')
    reflux_q3 = analysis_results.get('REFLUX_RATIO Q3')
    if isinstance(avg_reflux, (float, int)):
        summary_text += f"The column operated with a stable **Reflux Ratio (IQR Median)** of **{avg_reflux:.2f}**. The central 50% of operation fell between {reflux_q1:.2f} and {reflux_q3:.2f}. "
    
    naphthalene_loss_c01 = analysis_results.get('Naphthalene Loss in C-01 (%)')
    if isinstance(naphthalene_loss_c01, (float, int)):
        summary_text += f"The naphthalene loss in the C-01 column was calculated to be {naphthalene_loss_c01:.2f}%, which is "
        if naphthalene_loss_c01 > 2:
            summary_text += "**above the acceptable limit**. "
        else:
            summary_text += "within the acceptable limit. "

    naphthalene_c02_top = analysis_results.get('Naphthalene in C-02 Top Product (%)')
    if isinstance(naphthalene_c02_top, (float, int)):
        summary_text += f"Naphthalene concentration in the C-02 top product was found to be {naphthalene_c02_top:.2f}%. "
    
    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values are averages over the analysis period. Averages for Reflux Ratio and Heat Duties are calculated using the robust Median from the IQR method.")
    
    # Filter the results to only include the main KPI averages, not the Q1/Q3 values
    kpis_to_report = {k: v for k, v in analysis_results.items() if not k.endswith('Q1') and not k.endswith('Q3') and not k.endswith('Status')}

    for key, value in kpis_to_report.items():
        unit_key = KPI_UNIT_MAP.get(key)
        
        # Determine unit
        unit = ''
        if unit_key:
            # Look up unit using the mapped KPI key (e.g., 'REFLUX_RATIO')
            unit = TAG_UNITS.get(unit_key, '%')
        
        if isinstance(value, str):
            doc.add_paragraph(f"• **{key}**: {value}")
        elif isinstance(value, (float, int)):
            doc.add_paragraph(f"• **{key}**: {value:.2f} {unit}")

    # Section 3: Data Integrity and Outlier Analysis
    doc.add_heading('3. Data Integrity and Outlier Analysis', level=1)
    
    # 3.1 3-Sigma Outliers (Capped)
    doc.add_heading('3.1 Standard Outlier Capping ($\pm 3\sigma$)', level=2)
    doc.add_paragraph("Data for general process tags (flows, temperatures, DP) were capped to the $3\sigma$ limit.")
    if outliers_removed_3sigma:
        for tag, count in outliers_removed_3sigma.items():
            doc.add_paragraph(f"• **{tag}**: {count} outlier data points were capped.")
    else:
        doc.add_paragraph("No significant $3\sigma$ outliers were detected in general process tags.")

    # 3.2 IQR Outlier Filtering (Removed)
    doc.add_heading('3.2 Robust Average Calculation (IQR Method)', level=2)
    doc.add_paragraph("Outliers in critical metrics (Reflux Ratio and Heat Duties) were filtered using the Interquartile Range ($1.5 \times IQR$) method. The table below shows the Median (robust average) and the central 50% operating range ($Q1$ to $Q3$).")
    
    if iqr_outlier_summary:
        # Create a table for a cleaner display of the new summary data
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Median (Robust Average)'
        header_cells[2].text = 'Q1 (25th Percentile)'
        header_cells[3].text = 'Q3 (75th Percentile)'
        header_cells[4].text = 'Points Removed'
        header_cells[5].text = '% Removed'

        # Order the keys for better readability in the report
        ordered_keys = ['REFLUX_RATIO', 'REBOILER_HEAT_DUTY', 'CONDENSER_HEAT_DUTY']
        
        for tag in ordered_keys:
            if tag in iqr_outlier_summary:
                report = iqr_outlier_summary[tag]
                
                # Determine unit for display
                unit_display = f" ({TAG_UNITS.get(tag, '')})"
                if tag == 'REFLUX_RATIO':
                    unit_display = '' # Dimensionless

                row_cells = table.add_row().cells
                row_cells[0].text = tag
                row_cells[1].text = f"{report['Median_Value']:.2f}{unit_display}"
                row_cells[2].text = f"{report['Q1']:.2f}{unit_display}"
                row_cells[3].text = f"{report['Q3']:.2f}{unit_display}"
                row_cells[4].text = f"{report['Outlier_Count']:,}"
                row_cells[5].text = f"{report['Percentage_Removed']:.2f}%"
    else:
        doc.add_paragraph("No data available or no outliers detected for IQR metrics.")

    # Section 4: Material Balance Analysis
    doc.add_heading('4. Material Balance Analysis', level=1)
    
    doc.add_heading('4.1 C-00 and C-01 Material Balance', level=2)
    doc.add_paragraph(f"• C-00 Overall Material Balance Error: {analysis_results.get('C-00 Overall Material Balance Error (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"• Naphthalene Loss in C-01: {analysis_results.get('Naphthalene Loss in C-01 (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"    - {analysis_results.get('C-01 Naphthalene Loss Status', 'N/A')}")

    doc.add_heading('4.2 C-02 Material Balance', level=2)
    doc.add_paragraph(f"• C-02 Overall Material Balance Error: {analysis_results.get('C-02 Overall Material Balance Error (%)', 'N/A'):.2f}%")
    
    doc.add_heading('4.3 Component-wise Balance', level=2)
    doc.add_paragraph("Detailed component balance requires lab data on all input/output streams.")

    # Section 5: Performance Plots
    doc.add_heading('5. Performance Plots', level=1)
    
    plot_files = {
        '5.1 Naphthalene in Top Product vs. Reflux Ratio': output_naphthalene_vs_reflux_plot_path,
        '5.2 Naphthalene in Top Product vs. Reboiler Temperature': output_naphthalene_vs_reboiler_temp_plot_path,
        '5.3 Feed vs. Differential Pressure': output_feed_vs_dp_plot_path,
        '5.4 Temperature Profile': output_temp_plot_path,
        '5.5 Differential Pressure (DP)': output_dp_plot_path,
        '5.6 Daily Trends': output_trends_plot_path
    }
    
    for title, path in plot_files.items():
        doc.add_heading(title, level=2)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6))
        else:
            doc.add_paragraph(f"Plot file not found: {path}")

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

def main():
    """Main execution function."""
    # 1. Connect
    engine = connect_to_database()
    if engine is None:
        return

    # 2. Get Data
    scada_data, start_date, end_date = get_scada_data(engine)
    if scada_data is None:
        return

    # 3. Analyze and Clean
    # Now returns: analysis_results, scada_data (with cleaned values), 3-sigma outliers, IQR outlier summary
    analysis_results, scada_data, outliers_removed_3sigma, iqr_outlier_summary = perform_analysis(scada_data)
    
    if analysis_results:
        # 4. Generate Artifacts
        generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, outliers_removed_3sigma, iqr_outlier_summary, start_date, end_date)
        print("C-02 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()
