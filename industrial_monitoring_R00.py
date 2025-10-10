import numpy as np
import pandas as pd
from sklearn.ensemble import IsolationForest
from sklearn.preprocessing import StandardScaler
import joblib
import matplotlib.pyplot as plt
import os
import logging
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches
from prophet import Prophet
import io

# --- DYNAMIC PATH ADDITION ---
module_path = 'E:/GENERATING_DATA/ML_work'
if module_path not in sys.path:
    sys.path.append(module_path)
# --- END DYNAMIC PATH ADDITION ---

# Import the data retrieval and config
from data_retrieval_R00 import main_data_retrieval
from config_R00 import START_DATE, END_DATE, SCADA_TAGS_C00, SCADA_TAGS_C01, SCADA_TAGS_C02, SCADA_TAGS_C03, PM_TAGS, PM_THRESHOLDS

# --- LOGGING CONFIGURATION ---
# Correct logging setup to only show basic info to console
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

# File handler to capture all logs, including verbose ones
file_handler = logging.FileHandler("industrial_monitoring.log")
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

# Root logger configuration
logging.basicConfig(
    level=logging.INFO,
    handlers=[file_handler, console_handler]
)

SCADA_TAGS_BY_COLUMN = {
    'C-00': SCADA_TAGS_C00,
    'C-01': SCADA_TAGS_C01,
    'C-02': SCADA_TAGS_C02,
    'C-03': SCADA_TAGS_C03
}

# --- Utility Functions for Reporting ---
def create_word_report(column_id, model_name, pm_alert_msg, ad_alert_msg, pm_plot_buffer, ad_plot_buffer, start_date, end_date):
    logging.info(f"Generating Word report for {column_id} using {model_name} model...")
    document = Document()
    document.add_heading(f'Analysis Report for Column {column_id} ({model_name} Model)', level=1)
    
    document.add_paragraph(f"Report Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph(f"Data Analyzed From: {start_date} to {end_date}")
    document.add_heading('1. Predictive Maintenance Analysis', level=2)
    document.add_paragraph('This analysis predicts the future trend of a key sensor to forecast potential equipment failure.')
    document.add_paragraph(f'Status: {pm_alert_msg}')
    if pm_plot_buffer:
        document.add_picture(pm_plot_buffer, width=Inches(6))
    
    document.add_heading('2. Real-time Anomaly Detection Analysis', level=2)
    document.add_paragraph('This analysis identifies unusual sensor readings that deviate from normal operating behavior.')
    document.add_paragraph(f'Status: {ad_alert_msg}')
    if ad_plot_buffer:
        document.add_picture(ad_plot_buffer, width=Inches(6))
    
    report_filename = f'Analysis_Report_{column_id}_{model_name}.docx'
    document.save(report_filename)
    logging.info(f"Report for {column_id} saved to {report_filename}")

def generate_pm_plot(df, forecast_df, column_id, target_tag, model_name):
    plt.figure(figsize=(12, 6))
    plt.plot(df['ds'], df['y'], label='Historical Data', color='blue')
    plt.plot(forecast_df['ds'], forecast_df['yhat'], label='Forecast', color='red', linestyle='--')
    plt.title(f"Predictive Maintenance for {column_id} - {target_tag} ({model_name})")
    plt.xlabel("Date")
    plt.ylabel(f"Value of {target_tag}")
    plt.legend()
    plt.grid(True)
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return buf

def generate_ad_plot(scada_df, ad_cols, anomalies):
    if len(ad_cols) < 2:
        return None
    features_2d = scada_df.dropna(subset=ad_cols)[ad_cols[:2]].values
    plt.figure(figsize=(10, 8))
    plt.scatter(features_2d[:, 0], features_2d[:, 1], c=anomalies, cmap='viridis')
    plt.scatter(features_2d[anomalies == -1, 0], features_2d[anomalies == -1, 1], c='red', s=200, label='Anomaly')
    plt.title(f'Anomaly Detection using {ad_cols[0]} and {ad_cols[1]}')
    plt.xlabel(ad_cols[0])
    plt.ylabel(ad_cols[1])
    plt.legend()
    plt.grid(True)
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return buf

# --- Model Training and Saving ---
def train_and_save_models(scada_data, column_id, model_name):
    logging.info(f"--- Training and saving models for {column_id} using {model_name} model ---")
    if scada_data is None or scada_data.empty:
        logging.warning("No data available to train models. Skipping.")
        return False
    target_tag = PM_TAGS.get(column_id)
    pm_model_file = f'{column_id}_{model_name}_pm_model.pkl'
    if target_tag is None or target_tag not in scada_data.columns:
        logging.error(f"Target PM tag '{target_tag}' not found in data for {column_id}. Skipping PM model.")
        joblib.dump(None, pm_model_file)
    else:
        df_prophet = scada_data.rename(columns={'DATEANDTIME': 'ds', target_tag: 'y'})
        df_prophet = df_prophet[['ds', 'y']].dropna()
        if not df_prophet.empty:
            model_pm = Prophet()
            model_pm.fit(df_prophet)
            joblib.dump(model_pm, pm_model_file)
            joblib.dump(target_tag, f'{column_id}_pm_tag.pkl')
            logging.info(f"PM model for {column_id} trained and saved successfully.")
        else:
            logging.warning(f"PM data for {target_tag} is empty. Skipping PM model.")
            joblib.dump(None, pm_model_file)
    numerical_cols = [col for col in scada_data.columns if col not in ['DATEANDTIME', 'DateAndTime', 'dateandtime']]
    if len(numerical_cols) < 2:
        logging.warning(f"Not enough numerical tags for anomaly detection for {column_id}. Skipping AD model.")
        joblib.dump(None, f'{column_id}_ad_model.pkl')
        joblib.dump(None, f'{column_id}_ad_scaler.pkl')
        joblib.dump(None, f'{column_id}_ad_cols.pkl')
        return True
    features = scada_data.dropna(subset=numerical_cols)[numerical_cols].values
    scaler = StandardScaler()
    scaled_features = scaler.fit_transform(features)
    model_ad = IsolationForest(contamination=0.01, random_state=42)
    model_ad.fit(scaled_features)
    joblib.dump(model_ad, f'{column_id}_ad_model.pkl')
    joblib.dump(scaler, f'{column_id}_ad_scaler.pkl')
    joblib.dump(numerical_cols, f'{column_id}_ad_cols.pkl')
    logging.info(f"AD model for {column_id} trained and saved successfully.")
    return True

# --- Real-time Monitoring and Reporting ---
def monitor_real_time(column_id, scada_df, model_name):
    logging.info(f"--- Monitoring {column_id} in real-time using {model_name} model ---")
    try:
        pm_model = joblib.load(f'{column_id}_{model_name}_pm_model.pkl')
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
        pm_tag = joblib.load(f'{column_id}_pm_tag.pkl')
    except FileNotFoundError as e:
        logging.error(f"Models for {column_id} not found. Please train them first by running with 'train' mode.")
        logging.error(f"Error: {e}")
        return
    pm_alert_msg = "No immediate failure predicted within the next 50 days."
    pm_plot_buffer = None
    if pm_model and pm_tag and pm_tag in scada_df.columns:
        df_prophet = scada_df.rename(columns={'DATEANDTIME': 'ds', pm_tag: 'y'})
        df_prophet = df_prophet[['ds', 'y']].dropna()
        if not df_prophet.empty:
            future = pm_model.make_future_dataframe(periods=50)
            forecast = pm_model.predict(future)
            critical_threshold = PM_THRESHOLDS.get(pm_tag, None)
            if critical_threshold is not None:
                if any(forecast['yhat'] >= critical_threshold):
                    pm_alert_msg = f"ALERT: Failure predicted within the next 50 days based on {pm_tag}!"
            else:
                logging.warning(f"No critical threshold defined for {pm_tag}. Skipping PM alert.")
            pm_plot_buffer = generate_pm_plot(df_prophet, forecast, column_id, pm_tag, model_name)
        else:
            logging.warning(f"PM data for {pm_tag} is empty. Skipping PM analysis.")
    else:
        logging.warning(f"PM model or tag not available for {column_id}.")
    ad_alert_msg = "No significant anomalies detected."
    ad_plot_buffer = None
    if ad_model and ad_cols and len(ad_cols) >= 2:
        full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
        if not full_features_df.empty:
            scaled_features = scaler.transform(full_features_df.values)
            anomalies = ad_model.predict(scaled_features)
            if -1 in anomalies:
                ad_alert_msg = "ALERT: Anomaly detected in the data stream!"
            ad_plot_buffer = generate_ad_plot(full_features_df, ad_cols, anomalies)
        else:
            logging.warning("AD data is empty. Skipping AD analysis.")
    else:
        logging.warning(f"AD model not available or not enough tags for {column_id}.")
    create_word_report(column_id, model_name, pm_alert_msg, ad_alert_msg, pm_plot_buffer, ad_plot_buffer, START_DATE, END_DATE)

# --- Anomaly Analysis Function (Generates Ranked List File) ---
# --- Anomaly Analysis Function (Generates Ranked List File) ---
def analyze_anomalies(column_id, scada_df):
    logging.info(f"--- Starting Anomaly Ranking for {column_id} ---")
    try:
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
    except FileNotFoundError:
        logging.error(f"Models for {column_id} not found. Please train them first.")
        return
    if ad_model is None or ad_cols is None or len(ad_cols) < 2:
        logging.warning("Anomaly detection not configured for this column. Skipping.")
        return
    full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
    if full_features_df.empty:
        logging.warning("Data for anomaly analysis is empty. Skipping.")
        return
    scaled_features = scaler.transform(full_features_df.values)
    anomalies = ad_model.predict(scaled_features)
    anomaly_indices = np.where(anomalies == -1)[0]
    total_anomalies = len(anomaly_indices)
    if total_anomalies == 0:
        logging.info("No anomalies detected in the dataset.")
        return
    logging.info(f"Total Anomalies Detected: {total_anomalies}. A ranked list of instruments has been saved.")
    log_filename = f"Anomaly_Rank_List_{column_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    with open(log_filename, 'w') as log_file:
        log_file.write(f"Instrument Anomaly Ranking for Column {column_id}\n")
        log_file.write("="*40 + "\n\n")
        anomaly_data = full_features_df.loc[anomaly_indices]
        normal_data = full_features_df[anomalies == 1]
        if not normal_data.empty:
            normal_mean = normal_data.mean()
            normal_std = normal_data.std()
            anomaly_scores = {}
            for col in ad_cols:
                if normal_std[col] > 0:
                    deviations = np.abs((anomaly_data[col] - normal_mean[col]) / normal_std[col])
                    # New ranking logic: Sum the absolute deviations
                    total_deviation = deviations.sum()
                    anomaly_scores[col] = total_deviation
                else:
                    anomaly_scores[col] = 0
            sorted_anomaly_scores = sorted(anomaly_scores.items(), key=lambda item: item[1], reverse=True)
            log_file.write(f"Total Anomalies Detected: {total_anomalies}\n\n")
            log_file.write("Top Instruments with Highest Anomaly Contribution:\n")
            total_score = sum(anomaly_scores.values())
            for rank, (instrument, score) in enumerate(sorted_anomaly_scores, 1):
                if total_score > 0:
                    percentage = (score / total_score) * 100
                    log_file.write(f"{rank}. {instrument}: {score:.2f} total deviation ({percentage:.2f}% of total score)\n")
                else:
                    log_file.write(f"{rank}. {instrument}: {score:.2f} total deviation (0.00% of total score)\n")
        else:
            log_file.write("Could not generate ranking: No normal data points found to create a baseline.\n")
    logging.info(f"Ranked list of instruments saved to {log_filename}")
    
# --- Bulk Export to Excel ---
def export_all_anomalies_to_excel(column_id, scada_df):
    logging.info(f"--- Starting bulk export of anomalies for {column_id} ---")
    try:
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
    except FileNotFoundError:
        logging.error(f"Models for {column_id} not found. Please train them first.")
        return
    if ad_model is None or ad_cols is None or len(ad_cols) < 2:
        logging.warning("Anomaly detection not configured for this column. Skipping export.")
        return
    full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
    if full_features_df.empty:
        logging.warning("Data for anomaly export is empty. Skipping.")
        return
    scaled_features = scaler.transform(full_features_df.values)
    anomalies = ad_model.predict(scaled_features)
    anomaly_indices = np.where(anomalies == -1)[0]
    if len(anomaly_indices) == 0:
        logging.info(f"No anomalies detected for {column_id}.")
        return
    anomalous_data = full_features_df.iloc[anomaly_indices].copy()
    anomalous_data['Timestamp'] = scada_df.loc[anomalous_data.index]['DATEANDTIME']
    anomalous_data = anomalous_data[['Timestamp'] + ad_cols]
    filename = f"All_Anomalies_{column_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    try:
        anomalous_data.to_excel(filename, index=False)
        logging.info(f"Successfully exported {len(anomaly_indices)} anomalies to {filename}")
        logging.info("This file contains the full sensor data for each anomaly.")
    except Exception as e:
        logging.error(f"An error occurred while exporting to Excel: {e}")

# --- Main Execution Loop ---
def main(mode='monitor', model_name='prophet'):
    logging.info("="*60)
    logging.info(f"Starting program in '{mode}' mode with '{model_name}' model.")
    logging.info("="*60)
    for column_id, tags in SCADA_TAGS_BY_COLUMN.items():
        logging.info("="*60)
        scada_data, _ = main_data_retrieval(column_id, tags, START_DATE, END_DATE, 'GC_FILE', [])
        if scada_data is not None and not scada_data.empty:
            if mode == 'train':
                train_and_save_models(scada_data, column_id, model_name)
            elif mode == 'monitor':
                monitor_real_time(column_id, scada_data, model_name)
            elif mode == 'analyze':
                analyze_anomalies(column_id, scada_data)
            elif mode == 'bulk_export':
                export_all_anomalies_to_excel(column_id, scada_data)
            else:
                logging.error(f"Unknown mode: {mode}. Please use 'train', 'monitor', 'analyze', or 'bulk_export'.")
                break
        else:
            logging.warning(f"No SCADA data available for {column_id}. Skipping processing.")

# --- Main Entry Point ---
if __name__ == '__main__':
    if len(sys.argv) > 2:
        mode = sys.argv[1]
        model_name = sys.argv[2]
        main(mode=mode, model_name=model_name)
    else:
        main(mode='monitor', model_name='prophet')