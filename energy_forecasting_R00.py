import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
import joblib
from docx import Document
from docx.shared import Inches
from datetime import datetime
import sys
import pmdarima as pm
from sklearn.metrics import mean_squared_error
from sklearn.preprocessing import StandardScaler

# --- Define the path to the directory containing your modules ---
module_path = 'E:/GENERATING_DATA/ML_work'
if module_path not in sys.path:
    sys.path.append(module_path)

from data_retrieval_R00 import main_data_retrieval
from config_R00 import START_DATE, END_DATE, GC_FILE, SCADA_TAGS_C00, SCADA_TAGS_C01, SCADA_TAGS_C02, SCADA_TAGS_C03

output_plot_path = "energy_forecast_plot.png"
output_report_path = "Energy_Forecasting_Report.docx"
model_path = "energy_sarimax_model.pkl"
scaler_y_path = "y_scaler.pkl"
scaler_X_path = "X_scaler.pkl"

ENERGY_TAGS = [
    "HeatDuty_C-00_Reboiler",
    "HeatDuty_C-01_Reboiler",
    "HeatDuty_C-02_Reboiler",
    "HeatDuty_C-03_Reboiler"
]

# NEW: Define a list of exogenous variables (tags that influence energy)
EXOGENOUS_TAGS = [
    "FI-101", "FT-02", "FT-03", "FT-04", "FT-05", "FT-06", "FT-07", "FT-08", "FT-09", "FT-10", "FT-61", "FT-62", "FT-63", "FT-64",
    "TI-01", "TI-02", "TI-03", "TI-04", "TI-05", "TI-06", "TI-07", "TI-08", "TI-10", "TI-11", "TI-12", "TI-13", "TI-14", "TI-15", "TI-16", "TI-17", "TI-18", "TI-19", "TI-20", "TI-21", "TI-22", "TI-23", "TI-24", "TI-52", "TI-53", "TI-54", "TI-61", "TI-63", "TI-64", "TI-110", "TI-111", "TI-202", "TI-203", "TI-204", "TI-205", "TI-206", "TI-207", "TI-208", "TI-209", "TI-210", "TI-211", "TI-212", "TI-213", "TI-214", "TI-215", "TI-216"
]

# Global variables to store the scalers
y_scaler = None # Scaler for energy data (y)
X_scaler = None # Scaler for exogenous data (X)

def get_energy_data(start_date, end_date, sample_frequency='H'):
    """
    Retrieves and prepares energy consumption data and exogenous variables.
    """
    global y_scaler, X_scaler 
    print("\n--- Starting Data Retrieval ---")
    
    # NEW: Include both energy and exogenous tags in the retrieval
    all_tags = list(set(ENERGY_TAGS + EXOGENOUS_TAGS))
    
    scada_data, _ = main_data_retrieval(
        column_id=None,
        scada_tags=all_tags, 
        start_date=start_date, 
        end_date=end_date, 
        gc_file_path=GC_FILE, 
        gc_stream_ids=[]
    )

    if scada_data is None or scada_data.empty:
        print("🔴 ERROR: main_data_retrieval returned an empty or None DataFrame.")
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        
    found_energy_columns = [tag for tag in ENERGY_TAGS if tag in scada_data.columns]
    found_exog_columns = [tag for tag in EXOGENOUS_TAGS if tag in scada_data.columns]

    if not found_energy_columns:
        print("🔴 ERROR: No HeatDuty columns found in the retrieved data.")
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
        
    scada_data['TOTAL_ENERGY_KWH'] = scada_data[found_energy_columns].sum(axis=1)
    
    # Set index and resample
    energy_df = scada_data[['DATEANDTIME', 'TOTAL_ENERGY_KWH'] + found_exog_columns].copy()
    energy_df.set_index('DATEANDTIME', inplace=True)
    energy_df = energy_df.resample(sample_frequency).mean().dropna()

    y_data = energy_df[['TOTAL_ENERGY_KWH']]
    X_data = energy_df[found_exog_columns]

    # --- Use separate StandardScalers for y and X ---
    y_scaler = StandardScaler()
    scaled_y = y_scaler.fit_transform(y_data)
    
    X_scaler = StandardScaler()
    scaled_X = X_scaler.fit_transform(X_data)
    
    y_df_scaled = pd.DataFrame(scaled_y, index=y_data.index, columns=['Consumption_kWh'])
    X_df_scaled = pd.DataFrame(scaled_X, index=X_data.index, columns=found_exog_columns)

    # Split data for training and testing
    split_point = int(len(y_df_scaled) * 0.9)
    y_train = y_df_scaled.iloc[:split_point]
    y_test = y_df_scaled.iloc[split_point:]
    
    X_train = X_df_scaled.iloc[:split_point]
    X_test = X_df_scaled.iloc[split_point:]

    print("\n✅ Energy data and exogenous data preparation complete.")
    return y_train, y_test, X_train, X_test

def train_and_save_model(y_train, X_train):
    print("\n--- Training & Tuning SARIMAX model ---")
    if y_train.empty:
        print("🔴 ERROR: Cannot train on empty data.")
        return None
    try:
        model = pm.auto_arima(y=y_train, # Pass the energy data
                              X=X_train, # Pass the exogenous data
                              start_p=1, start_q=1,
                              test='adf',       
                              max_p=5, max_q=5,
                              max_d=2,          
                              m=24,             
                              start_P=0, seasonal=True,
                              D=1,              
                              trace=True,
                              error_action='ignore',  
                              suppress_warnings=True,
                              stepwise=True)
        
        # Save the model and both scalers
        joblib.dump(model, model_path)
        joblib.dump(y_scaler, scaler_y_path)
        joblib.dump(X_scaler, scaler_X_path)

        print("✅ SARIMAX model and scalers saved successfully.")
        return model
    except Exception as e:
        print(f"🔴 ERROR: Model training failed. Reason: {e}")
        return None

def run_forecast_and_report(model_fit, y_train, y_test, X_test, forecast_steps, start_date_str, end_date_str):
    print(f"\n--- Starting Forecast for {forecast_steps} hours ---")
    
    # Load scalers
    y_scaler_loaded = joblib.load(scaler_y_path)
    X_scaler_loaded = joblib.load(scaler_X_path)

    # In-sample forecast for validation
    forecast_test, conf_int = model_fit.predict(n_periods=len(y_test), X=X_test, return_conf_int=True)
    
    # Inverse transform the forecast back to original units (kWh)
    forecast_test_original = y_scaler_loaded.inverse_transform(forecast_test.values.reshape(-1, 1)).flatten()
    forecast_test_series = pd.Series(forecast_test_original, index=y_test.index)

    # Inverse transform the test data for RMSE calculation
    y_test_original = y_scaler_loaded.inverse_transform(y_test[['Consumption_kWh']]).flatten()
    rmse = np.sqrt(mean_squared_error(y_test_original, forecast_test_original))
    print(f"✅ Model validation complete. RMSE on test data: {rmse:.2f}")

    # NEW: Out-of-sample forecast for the future
    # NOTE: For a true future forecast, you would need to forecast the exogenous variables (X_future).
    # For now, we'll create a dummy future X_future based on the last known values.
    # In a real-world scenario, you would need a separate model to forecast these.
    last_known_X_values = X_test.iloc[-1].values.reshape(1, -1)
    X_future_scaled = np.tile(last_known_X_values, (forecast_steps, 1))
    
    forecast_future, conf_int_future = model_fit.predict(n_periods=forecast_steps, X=X_future_scaled, return_conf_int=True)

    # Inverse transform the future forecast
    forecast_future_original = y_scaler_loaded.inverse_transform(forecast_future.values.reshape(-1, 1)).flatten()
    forecast_index_future = pd.date_range(start=y_test.index[-1] + pd.Timedelta(hours=1), periods=forecast_steps, freq='H')
    forecast_future_series = pd.Series(forecast_future_original, index=forecast_index_future)

    print("\n✅ Forecast complete. Generating report...")
    
    # Inverse transform the training and test data for plotting
    y_train_original = y_scaler_loaded.inverse_transform(y_train[['Consumption_kWh']]).flatten()
    plot_forecast(y_train_original, y_test_original, forecast_test_series, forecast_future_series, y_test.index)
    generate_word_report(model_fit, forecast_future_series, rmse, output_plot_path, start_date_str, end_date_str)

def generate_word_report(model_fit, forecast_series, rmse, plot_path, start_date_str, end_date_str):
    """Creates a detailed Word document report of the energy forecast."""
    doc = Document()
    doc.add_heading('Energy Consumption Forecast Report', 0)
    
    doc.add_paragraph(f"Report Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Analysis Period: {start_date_str} to {end_date_str}")
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph("This report presents a forecast of future energy flow using a time-series model. The forecast is based on historical data patterns and key operational variables.")
    doc.add_heading('2. Model and Forecast Details', level=1)
    doc.add_paragraph(f"The forecast was generated using a **SARIMAX** model, which accounts for external factors.")
    if hasattr(model_fit, 'aic'):
        doc.add_paragraph(f"The model's AIC (Akaike Information Criterion) is: **{model_fit.aic():.2f}**")
    doc.add_paragraph(f"Model validation was performed on a test set, yielding an RMSE of: **{rmse:.2f}** kWh.")
    
    doc.add_heading('3. Key Forecasts (Next 5 Days)', level=1)
    doc.add_paragraph("The table below shows the key forecast values:")
    
    forecast_df = forecast_series.to_frame(name='Forecasted_Consumption_kWh')
    forecast_df.index.name = 'Date'
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Forecasted Energy (kWh)'
    
    for date, consumption in forecast_df.head(120).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(date.strftime('%Y-%m-%d %H:%M'))
        row_cells[1].text = f"{consumption.values[0]:.2f}"
    
    doc.add_heading('4. Visual Analysis', level=1)
    doc.add_paragraph("The chart below shows the historical energy flow and the future forecast.")
    if os.path.exists(plot_path):
        doc.add_picture(plot_path, width=Inches(6))
    
    doc.add_heading('5. Proactive Recommendation', level=1)
    if not forecast_series.empty:
        avg_forecast_consumption = forecast_series.mean()
        recommendation = ""
        if avg_forecast_consumption > 0:
            recommendation = "The forecast indicates a period of energy consumption by the system. It is recommended to review operational schedules and implement energy-saving measures to manage costs proactively."
        else:
            recommendation = "The forecast indicates the system is generating energy. Continue monitoring the process, but no immediate action is required based on the forecast."
        doc.add_paragraph(recommendation)
    else:
        doc.add_paragraph("No forecast data was available to provide a recommendation.")
    
    doc.save(output_report_path)
    print(f"\n✅ Energy forecast report generated successfully at {output_report_path}")

def plot_forecast(y_train_original, y_test_original, forecast_test_series, forecast_future_series, y_test_index):
    """Plots the historical data, test data, and the forecast."""
    plt.figure(figsize=(12, 6))
    plt.plot(y_train_original, label='Historical (Train) Data', color='blue')
    plt.plot(y_test_original, label='Historical (Test) Data', color='green')
    plt.plot(y_test_index, forecast_test_series.values, label='In-Sample Forecast', color='orange', linestyle='--')
    plt.plot(forecast_future_series.index, forecast_future_series.values, label='Future Forecast', color='red', linestyle='--')
    
    plt.title('Energy Flow Forecast')
    plt.xlabel('Date')
    plt.ylabel('Energy (kWh)')
    plt.legend()
    plt.grid(True)
    
    plt.savefig(output_plot_path)
    plt.close()
    print(f"✅ Forecast plot saved to {output_plot_path}")

# --- Main logic to run the script ---
def main(mode='train'):
    if mode == 'train':
        print("\n--- Running in TRAINING mode (SARIMAX) ---")
        y_train, y_test, X_train, X_test = get_energy_data(START_DATE, END_DATE)
        if not y_train.empty:
            model = train_and_save_model(y_train, X_train)
        else:
            print("🔴 TRAINING FAILED: Data retrieval returned empty data.")
            return

    elif mode == 'monitor':
        print("\n--- Running in MONITORING mode (SARIMAX) ---")
        try:
            model = joblib.load(model_path)
            y_scaler_loaded = joblib.load(scaler_y_path)
            X_scaler_loaded = joblib.load(scaler_X_path)
            print("✅ Model and scalers loaded successfully.")
        except FileNotFoundError:
            print(f"🔴 ERROR: Model or scaler files not found. Please run in 'train' mode first.")
            return
        
        y_train, y_test, X_train, X_test = get_energy_data(START_DATE, END_DATE)

        run_forecast_and_report(model, y_train, y_test, X_test, forecast_steps=120, 
                                start_date_str=START_DATE, end_date_str=END_DATE)

# --- Main file logic to run the script ---
if __name__ == "__main__":
    main(mode='train')
    main(mode='monitor')