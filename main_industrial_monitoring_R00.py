import numpy as np
import pandas as pd
from sklearn.ensemble import IsolationForest, RandomForestRegressor
from sklearn.preprocessing import StandardScaler
import joblib
import matplotlib.pyplot as plt
import os
import logging
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches
from prophet import Prophet

# --- DYNAMIC PATH ADDITION ---
module_path = 'E:/GENERATING_DATA/ML_work'
if module_path not in sys.path:
    sys.path.append(module_path)
# --- END DYNAMIC PATH ADDITION ---

# Import the data retrieval and config
from data_retrieval_R00 import main_data_retrieval
from config_R00 import START_DATE, END_DATE, GC_FILE, SCADA_TAGS_C00, SCADA_TAGS_C01, SCADA_TAGS_C02, SCADA_TAGS_C03, PM_TAGS, PM_THRESHOLDS

# --- LOGGING CONFIGURATION ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("industrial_monitoring.log"),
        logging.StreamHandler(sys.stdout)
    ]
)

SCADA_TAGS_BY_COLUMN = {
    'C-00': SCADA_TAGS_C00,
    'C-01': SCADA_TAGS_C01,
    'C-02': SCADA_TAGS_C02,
    'C-03': SCADA_TAGS_C03
}

# --- Utility Functions for Reporting ---

def create_word_report(column_id, pm_alert_msg, ad_alert_msg, pm_plot_path, ad_plot_path, start_date, end_date):
    logging.info(f"Generating Word report for {column_id}...")
    document = Document()
    document.add_heading(f'Analysis Report for Column {column_id}', level=1)
    
    document.add_paragraph(f"Report Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph(f"Data Analyzed From: {start_date} to {end_date}")

    document.add_heading('1. Predictive Maintenance Analysis', level=2)
    document.add_paragraph('This analysis predicts the future trend of a key sensor to forecast potential equipment failure.')
    document.add_paragraph(f'Status: {pm_alert_msg}')
    if os.path.exists(pm_plot_path):
        document.add_picture(pm_plot_path, width=Inches(6))
    
    document.add_heading('2. Real-time Anomaly Detection Analysis', level=2)
    document.add_paragraph('This analysis identifies unusual sensor readings that deviate from normal operating behavior.')
    document.add_paragraph(f'Status: {ad_alert_msg}')
    if os.path.exists(ad_plot_path):
        document.add_picture(ad_plot_path, width=Inches(6))
    
    report_filename = f'Analysis_Report_{column_id}.docx'
    document.save(report_filename)
    logging.info(f"Report for {column_id} saved to {report_filename}")

def generate_pm_plot(df, forecast_df, column_id, target_tag):
    pm_plot_path = f"{column_id}_predictive_maintenance_plot.png"
    plt.figure(figsize=(12, 6))
    
    plt.plot(df['ds'], df['y'], label='Historical Data', color='blue')
    plt.plot(forecast_df['ds'], forecast_df['yhat'], label='Forecast', color='red', linestyle='--')
    
    plt.title(f"Predictive Maintenance for {column_id} - {target_tag}")
    plt.xlabel("Date")
    plt.ylabel(f"Value of {target_tag}")
    plt.legend()
    plt.grid(True)
    plt.savefig(pm_plot_path)
    plt.close()
    logging.info(f"PM plot saved to {pm_plot_path}")
    return pm_plot_path

# --- Model Training and Saving ---

def train_and_save_models(scada_data, column_id):
    logging.info(f"--- Training and saving models for {column_id} ---")
    if scada_data is None or scada_data.empty:
        logging.warning("No data available to train models. Skipping.")
        return False

    # Get the specific PM tag for this column from the config
    target_tag = PM_TAGS.get(column_id)
    if target_tag is None or target_tag not in scada_data.columns:
        logging.error(f"Target PM tag '{target_tag}' not found in data for {column_id}. Skipping PM model.")
        joblib.dump(None, f'{column_id}_pm_model.pkl')
    else:
        # Prepare data for Prophet
        df_prophet = scada_data.rename(columns={'DATEANDTIME': 'ds', target_tag: 'y'})
        df_prophet = df_prophet[['ds', 'y']].dropna()

        if not df_prophet.empty:
            model_pm = Prophet()
            model_pm.fit(df_prophet)
            joblib.dump(model_pm, f'{column_id}_pm_model.pkl')
            joblib.dump(target_tag, f'{column_id}_pm_tag.pkl')
            logging.info(f"PM model for {column_id} trained and saved successfully.")
        else:
            logging.warning(f"PM data for {target_tag} is empty. Skipping PM model.")
            joblib.dump(None, f'{column_id}_pm_model.pkl')

    numerical_cols = [col for col in scada_data.columns if col not in ['DATEANDTIME', 'DateAndTime', 'dateandtime']]
    if len(numerical_cols) < 2:
        logging.warning(f"Not enough numerical tags for anomaly detection for {column_id}. Skipping AD model.")
        joblib.dump(None, f'{column_id}_ad_model.pkl')
        joblib.dump(None, f'{column_id}_ad_scaler.pkl')
        joblib.dump(None, f'{column_id}_ad_cols.pkl')
        return True
    
    features = scada_data.dropna(subset=numerical_cols)[numerical_cols].values
    scaler = StandardScaler()
    scaled_features = scaler.fit_transform(features)
    model_ad = IsolationForest(contamination='auto', random_state=42)
    model_ad.fit(scaled_features)

    joblib.dump(model_ad, f'{column_id}_ad_model.pkl')
    joblib.dump(scaler, f'{column_id}_ad_scaler.pkl')
    joblib.dump(numerical_cols, f'{column_id}_ad_cols.pkl')
    
    logging.info(f"AD model for {column_id} trained and saved successfully.")
    return True

# --- Real-time Monitoring and Reporting ---

def monitor_real_time(column_id, scada_df):
    logging.info(f"--- Monitoring {column_id} in real-time ---")
    
    try:
        pm_model = joblib.load(f'{column_id}_pm_model.pkl')
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
        pm_tag = joblib.load(f'{column_id}_pm_tag.pkl')
    except FileNotFoundError:
        logging.error(f"Models for {column_id} not found. Please train them first by running with 'train' mode.")
        return

    pm_alert_msg = "No immediate failure predicted within the next 50 days."
    pm_plot_path = "n/a"
    
    if pm_model and pm_tag and pm_tag in scada_df.columns:
        df_prophet = scada_df.rename(columns={'DATEANDTIME': 'ds', pm_tag: 'y'})
        df_prophet = df_prophet[['ds', 'y']].dropna()
        
        if not df_prophet.empty:
            future = pm_model.make_future_dataframe(periods=50)
            forecast = pm_model.predict(future)
            
            critical_threshold = PM_THRESHOLDS.get(pm_tag, None)
            if critical_threshold is not None:
                if any(forecast['yhat'] >= critical_threshold):
                    pm_alert_msg = f"ALERT: Failure predicted within the next 50 days based on {pm_tag}!"
            else:
                logging.warning(f"No critical threshold defined for {pm_tag}. Skipping PM alert.")
            
            pm_plot_path = generate_pm_plot(df_prophet, forecast, column_id, pm_tag)
        else:
            logging.warning(f"PM data for {pm_tag} is empty. Skipping PM analysis.")
    else:
        logging.warning(f"PM model or tag not available for {column_id}.")
        
    ad_alert_msg = "No significant anomalies detected."
    ad_plot_path = "n/a"
    
    if ad_model and ad_cols and len(ad_cols) >= 2:
        full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
        if not full_features_df.empty:
            scaled_features = scaler.transform(full_features_df.values)
            anomalies = ad_model.predict(scaled_features)
            
            features_2d = full_features_df[ad_cols[:2]].values
            
            ad_plot_path = f"{column_id}_anomaly_detection_plot.png"
            plt.figure(figsize=(10, 8))
            plt.scatter(features_2d[:, 0], features_2d[:, 1], c=anomalies, cmap='viridis')
            plt.scatter(features_2d[anomalies == -1, 0], features_2d[anomalies == -1, 1], c='red', s=200, label='Anomaly')
            plt.title(f'Anomaly Detection for {column_id} using {ad_cols[0]} and {ad_cols[1]}')
            plt.xlabel(ad_cols[0])
            plt.ylabel(ad_cols[1])
            plt.legend()
            plt.grid(True)
            plt.savefig(ad_plot_path)
            plt.close()

            if -1 in anomalies:
                ad_alert_msg = "ALERT: Anomaly detected in the data stream!"
        else:
            logging.warning("AD data is empty. Skipping AD analysis.")
    else:
        logging.warning(f"AD model not available or not enough tags for {column_id}.")

    create_word_report(column_id, pm_alert_msg, ad_alert_msg, pm_plot_path, ad_plot_path, START_DATE, END_DATE)


# --- New Anomaly Analysis Function ---

def analyze_anomalies(column_id, scada_df):
    logging.info(f"--- Starting Detailed Anomaly Analysis for {column_id} ---")
    
    try:
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
    except FileNotFoundError:
        logging.error(f"Models for {column_id} not found. Please train them first.")
        return

    if ad_model is None or ad_cols is None or len(ad_cols) < 2:
        logging.warning("Anomaly detection not configured for this column. Skipping.")
        return

    full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
    
    if full_features_df.empty:
        logging.warning("Data for anomaly analysis is empty. Skipping.")
        return

    scaled_features = scaler.transform(full_features_df.values)
    anomalies = ad_model.predict(scaled_features)
    
    anomaly_indices = np.where(anomalies == -1)[0]
    
    if len(anomaly_indices) == 0:
        logging.info("No anomalies detected in the dataset.")
        return

    logging.info(f"Total Anomalies Detected: {len(anomaly_indices)}")
    
    log_filename = f"Anomaly_Log_{column_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    with open(log_filename, 'w') as log_file:
        log_file.write(f"Anomaly Log for Column {column_id}\n")
        log_file.write("="*40 + "\n\n")

        for i, idx in enumerate(anomaly_indices):
            anomaly_data = full_features_df.iloc[idx]
            timestamp = scada_df.loc[anomaly_data.name, 'DATEANDTIME']
            
            log_file.write(f"Anomaly #{i+1} at {timestamp}\n")
            log_file.write("-" * 20 + "\n")

            for col in ad_cols:
                log_file.write(f"{col}: {anomaly_data[col]}\n")
            log_file.write("\n")

            logging.info(f"Logged anomaly at: {timestamp}")
            
    logging.info(f"Detailed anomaly log saved to {log_filename}")

# --- Bulk Export to Excel ---

def export_all_anomalies_to_excel(column_id, scada_df):
    logging.info(f"--- Starting bulk export of anomalies for {column_id} ---")

    try:
        ad_model = joblib.load(f'{column_id}_ad_model.pkl')
        scaler = joblib.load(f'{column_id}_ad_scaler.pkl')
        ad_cols = joblib.load(f'{column_id}_ad_cols.pkl')
    except FileNotFoundError:
        logging.error(f"Models for {column_id} not found. Please train them first.")
        return

    if ad_model is None or ad_cols is None or len(ad_cols) < 2:
        logging.warning("Anomaly detection not configured for this column. Skipping export.")
        return

    full_features_df = scada_df.dropna(subset=ad_cols)[ad_cols]
    
    if full_features_df.empty:
        logging.warning("Data for anomaly export is empty. Skipping.")
        return
        
    scaled_features = scaler.transform(full_features_df.values)
    anomalies = ad_model.predict(scaled_features)
    
    anomaly_indices = np.where(anomalies == -1)[0]
    
    if len(anomaly_indices) == 0:
        logging.info(f"No anomalies detected for {column_id}.")
        return

    anomalous_data = full_features_df.iloc[anomaly_indices].copy()
    
    anomalous_data['Timestamp'] = scada_df.loc[anomalous_data.index]['DATEANDTIME']
    
    anomalous_data = anomalous_data[['Timestamp'] + ad_cols]

    filename = f"All_Anomalies_{column_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    
    try:
        anomalous_data.to_excel(filename, index=False)
        logging.info(f"Successfully exported {len(anomaly_indices)} anomalies to {filename}")
        logging.info("This file contains the full sensor data for each anomaly.")
    except Exception as e:
        logging.error(f"An error occurred while exporting to Excel: {e}")

# --- Main Execution Loop ---

def main(mode='monitor'):
    logging.info("="*60)
    logging.info(f"Starting program in '{mode}' mode.")
    logging.info("="*60)

    for column_id, tags in SCADA_TAGS_BY_COLUMN.items():
        logging.info("="*60)
        scada_data, _ = main_data_retrieval(column_id, tags, START_DATE, END_DATE, GC_FILE, [])
        if scada_data is not None and not scada_data.empty:
            if mode == 'train':
                train_and_save_models(scada_data, column_id)
            elif mode == 'monitor':
                monitor_real_time(column_id, scada_data)
            elif mode == 'analyze':
                analyze_anomalies(column_id, scada_data)
            else:
                logging.error(f"Unknown mode: {mode}. Please use 'train', 'monitor', or 'analyze'.")
                break
        else:
            logging.warning(f"No SCADA data available for {column_id}. Skipping processing.")

    # Special handling for bulk_export mode as it requires a specific column ID
    if mode == 'bulk_export':
        if len(sys.argv) < 3:
            logging.error("Error: Please provide a column ID for bulk export.")
            logging.info("Example: python main.py bulk_export C-00")
            return
        column_to_export = sys.argv[2]
        tags_to_get = SCADA_TAGS_BY_COLUMN.get(column_to_export)
        
        if tags_to_get is None:
            logging.error(f"Error: Unknown column ID '{column_to_export}'.")
            return
            
        scada_data, _ = main_data_retrieval(column_to_export, tags_to_get, START_DATE, END_DATE, GC_FILE, [])
        if scada_data is not None and not scada_data.empty:
            export_all_anomalies_to_excel(column_to_export, scada_data)
        else:
            logging.warning(f"No data available for {column_to_export}.")
    elif mode not in ['train', 'monitor', 'analyze']:
        logging.error(f"Unknown mode: {mode}. Please use 'train', 'monitor', 'analyze', or 'bulk_export'.")

if __name__ == '__main__':
    if len(sys.argv) > 1:
        main(mode=sys.argv[1])
    else:
        main(mode='monitor')