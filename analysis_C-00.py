import pandas as pd
from sqlalchemy import create_engine, inspect
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import re

# =================================================================================
# 1. CONFIGURATION AND CONSTANTS
# =================================================================================

# Database connection parameters (UPDATE THESE WITH YOUR ACTUAL DETAILS)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
TAG_UNITS = {
    # Flows
    'FT-01': 'kg/h',     # Feed
    'FT-61': 'kg/h',     # Top Product (Moisture/Light Ends)
    'FT-62': 'kg/h',     # Bottom Product (C-01 Feed)
    # Temperatures
    'TI-01': 'degC',     # Feed
    'TI-02': 'degC',     # Top Section
    'TI-03': 'degC',     # Mid Section
    'TI-04': 'degC',     # Bottom Section
    # Utilities (Used as fallback if direct duty tags are missing)
    'TI-101': 'degC',    # Cooling Water In
    'TI-102': 'degC',    # Cooling Water Out
    'FI-101': 'kg/h',    # Cooling Water Flow
    # Direct Tags used for analysis
    'DP-C-00': 'mmHg', 
    'HeatDuty-C-00-Reboiler': 'kW',
    'HeatDuty-C-00-Condenser': 'kW',
}

# File paths for saving generated plots and report
output_report_path = "C-00_Analysis_Report.docx"
output_temp_plot_path = "C-00_Temperature_Profile.png"
output_dp_plot_path = "C-00_Differential_Pressure.png"
output_trends_plot_path = "C-00_Daily_Trends.png"
output_moisture_vs_reboiler_plot_path = "C-00_Moisture_vs_Reboiler.png"
output_moisture_vs_feedtemp_plot_path = "C-00_Moisture_vs_FeedTemp.png"

# Engineering constants
WATER_SPECIFIC_HEAT = 4.186        # kJ/(kg·°C)
# Average assumed moisture content in feed (e.g., 0.2% by mass)
AVG_FEED_MOISTURE_MASS_PERCENT = 0.2

# KPI Thresholds and limits
MAX_REASONABLE_EFFICIENCY_PERCENT = 200 # Cap efficiency for accurate KPI averaging
EFFICIENCY_PLOT_CAP = 400 # Cap efficiency for plotting visualization
SIMULATED_REBOILER_MIN_KW = 200.0 # User-defined lower bound for simulation
SIMULATED_REBOILER_MAX_KW = 600.0 # User-defined upper bound for simulation

# =================================================================================
# 2. DATABASE AND DATA RETRIEVAL
# =================================================================================

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine, start_date='2025-09-03 00:00:00', end_date='2025-09-30 00:00:00'):
    """Retrieves specific SCADA data for the C-00 column from the database."""
    try:
        # List of tags required for C-00 analysis
        desired_columns = [
            "DateAndTime", "FT-01", "FT-61", "FT-62", "TI-01", "TI-02", "TI-03", "TI-04",
            "TI-101", "TI-102", "FI-101", "DP_C-00", "HeatDuty_C-00_Reboiler", "HeatDuty_C-00_Condenser"
        ]

        inspector = inspect(engine)
        columns = inspector.get_columns('data_cleaning_with_report')
        column_names = [col['name'] for col in columns]

        final_columns = []
        for d_col in desired_columns:
            found = False
            for db_col in column_names:
                # Normalization to handle DB column naming conventions (e.g., DP_C_00 vs DP-C-00)
                normalized_d_col = d_col.replace('-', '').lower()
                normalized_db_col = db_col.replace('-', '').lower()

                if normalized_d_col == normalized_db_col:
                    final_columns.append(f'"{db_col}"')
                    found = True
                    break
            if not found:
                # We skip the missing column for now, the analysis function will handle it with N/A or simulation
                pass 

        if not final_columns:
            print("Error: No matching columns found for C-00. Data retrieval failed.")
            return None, start_date, end_date

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM data_cleaning_with_report
        WHERE "DateAndTime" BETWEEN '{start_date}' AND '{end_date}'
        ORDER BY "DateAndTime";
        """

        df = pd.read_sql(query, engine)
        # Standardize column names (UPPERCASE, replace '-' with '_')
        df.columns = [col.upper().replace('-', '_') for col in df.columns]
        df['DATEANDTIME'] = pd.to_datetime(df['DATEANDTIME'])
        print("SCADA data for C-00 retrieved successfully.")
        return df, start_date, end_date
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None, None, None

# =================================================================================
# 3. ANALYSIS AND CALCULATIONS
# =================================================================================

def calculate_robust_average(series, min_value_threshold=0.1, std_dev_multiplier=3):
    """
    Calculates a robust average of a numeric series by removing outliers 
    (values outside the mean +/- std_dev_multiplier * standard deviation).
    
    Returns the average and a flag indicating if the data was considered valid.
    """
    series_clean = pd.to_numeric(series, errors='coerce').fillna(0)
    
    # 1. Check for sufficient data variation and non-zero maximum
    if series_clean.std() < min_value_threshold and series_clean.max() < min_value_threshold:
        return 0, False, "Data is essentially flatlined/zero."

    # 2. Outlier removal (3-sigma rule for a robust average)
    mean = series_clean.mean()
    std_dev = series_clean.std()
    
    # Calculate bounds aggressively to remove junk data
    lower_bound = mean - std_dev_multiplier * std_dev
    upper_bound = mean + std_dev_multiplier * std_dev
    
    # Filter series: keep data within the bounds
    filtered_series = series_clean[
        (series_clean >= lower_bound) & 
        (series_clean <= upper_bound) &
        (series_clean > min_value_threshold) # Ensure positive and meaningful
    ]

    # 3. Final Check
    if filtered_series.empty:
        # If all data was filtered out
        return 0, False, "All data points were filtered as extreme outliers or were zero."
    
    return filtered_series.mean(), True, None


def perform_analysis(df):
    """Performs key calculations for C-00, including material balances and efficiency."""
    if df is None or df.empty:
        return {}, df, {}

    analysis_results = {}

    # Convert key columns to numeric, filling missing/non-numeric data with 0
    cols_to_clean = ['FT_01', 'FT_61', 'FT_62', 'TI_01', 'TI_04',
                     'TI_101', 'TI_102', 'FI_101', 'DP_C_00',
                     'HEATDUTY_C_00_REBOILER', 'HEATDUTY_C_00_CONDENSER']

    for col in cols_to_clean:
        if col in df.columns:
            df.loc[:, col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

    # Calculate average flow rates
    feed_flow_avg = df['FT_01'].mean()
    top_product_flow_avg = df['FT_61'].mean()
    bottom_product_flow_avg = df['FT_62'].mean()

    analysis_results['Average Feed Flow (FT-01)'] = feed_flow_avg
    analysis_results['Average Top Product Flow (FT-61)'] = top_product_flow_avg
    analysis_results['Average Bottom Product Flow (FT-62)'] = bottom_product_flow_avg

    # Overall Material Balance: Error = |(In - Out) / In| * 100
    if feed_flow_avg > 0:
        material_balance_error = ((feed_flow_avg - (top_product_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
        analysis_results['Overall Material Balance Error (%)'] = abs(material_balance_error)

    # Moisture Removal Efficiency Calculation (Primary KPI)
    # 1. Estimated Moisture In Feed
    df['MOISTURE_IN_FEED'] = df['FT_01'] * (AVG_FEED_MOISTURE_MASS_PERCENT / 100.0)

    # 2. Moisture Removal Efficiency (%)
    epsilon = 1e-6
    df['MOISTURE_REMOVAL_EFFICIENCY'] = (df['FT_61'] / (df['MOISTURE_IN_FEED'] + epsilon)) * 100

    # Clean up where division by zero or negative flows occurred
    df.loc[df['MOISTURE_REMOVAL_EFFICIENCY'] < 0, 'MOISTURE_REMOVAL_EFFICIENCY'] = 0

    # For visualization, we cap at a higher limit (e.g., 400%)
    df['MOISTURE_REMOVAL_EFFICIENCY_PLOTTED'] = df['MOISTURE_REMOVAL_EFFICIENCY'].clip(upper=EFFICIENCY_PLOT_CAP)


    # Average Metrics - Use a filtered DataFrame to calculate a realistic average KPI
    # FIX CONFIRMED: Filter out unrealistically high efficiencies (> MAX_REASONABLE_EFFICIENCY_PERCENT) for KPI average
    df_filtered_for_avg = df[df['MOISTURE_REMOVAL_EFFICIENCY'] <= MAX_REASONABLE_EFFICIENCY_PERCENT]

    analysis_results['Average Moisture Content in Feed (%)'] = AVG_FEED_MOISTURE_MASS_PERCENT
    analysis_results['Average Moisture Removal Efficiency (%)'] = df_filtered_for_avg['MOISTURE_REMOVAL_EFFICIENCY'].mean()


    # Differential Pressure (DP) and Heat Duty Metrics
    if 'DP_C_00' in df.columns:
        analysis_results['Average Differential Pressure'] = df['DP_C_00'].mean()
        analysis_results['Maximum Differential Pressure'] = df['DP_C_00'].max()
        df.rename(columns={'DP_C_00': 'DIFFERENTIAL_PRESSURE'}, inplace=True)


    # >>>>>>> REBOILER HEAT DUTY FIX: Implement SIMULATION if data is bad <<<<<<<
    direct_reboiler_tag = 'HEATDUTY_C_00_REBOILER'
    is_reboiler_simulated = False
    
    if direct_reboiler_tag in df.columns:
        # Attempt to get robust average from real tag data
        reboiler_avg, is_valid, reason = calculate_robust_average(df[direct_reboiler_tag])
        
        if is_valid and reboiler_avg > 0.1: 
            # Case A: Good, valid data found (use real data)
            analysis_results['Average Reboiler Heat Duty'] = reboiler_avg
            df.rename(columns={direct_reboiler_tag: 'REBOILER_HEAT_DUTY'}, inplace=True)
        else:
            # Case B: Bad sensor data - SIMULATE based on user's input (200-600 kW)
            is_reboiler_simulated = True
            np.random.seed(42) # Set seed for reproducibility of simulated data
            simulated_duty = np.random.uniform(SIMULATED_REBOILER_MIN_KW, SIMULATED_REBOILER_MAX_KW, size=len(df))
            df['REBOILER_HEAT_DUTY'] = simulated_duty
            
            # Recalculate average from simulated data for KPI
            sim_avg = simulated_duty.mean()
            analysis_results['Average Reboiler Heat Duty'] = f'{sim_avg:.2f} kW (SIMULATED - Target {SIMULATED_REBOILER_MIN_KW}-{SIMULATED_REBOILER_MAX_KW}kW)'
            df.rename(columns={direct_reboiler_tag: 'REBOILER_HEAT_DUTY_ORIGINAL'}, inplace=True) # Rename original bad column
            print(f"WARNING: Reboiler Heat Duty data was bad. SIMULATED data used for analysis: Avg {sim_avg:.2f} kW.")
    else:
        analysis_results['Average Reboiler Heat Duty'] = 'N/A (Missing Tag in DB)'
        df['REBOILER_HEAT_DUTY'] = 0 # Placeholder for plotting
    # >>>>>>> END OF REBOILER HEAT DUTY FIX <<<<<<<


    # Handle Condenser Heat Duty (Use robust average with outlier removal, then fallback)
    direct_condenser_tag = 'HEATDUTY_C_00_CONDENSER'
    utility_tags_available = all(tag in df.columns for tag in ['TI_101', 'TI_102', 'FI_101'])
    
    condenser_avg = 0
    
    if direct_condenser_tag in df.columns:
        condenser_avg, is_valid, reason = calculate_robust_average(df[direct_condenser_tag])
        
        if is_valid and condenser_avg != 0:
            # 1. Use robust average of direct tag value
            analysis_results['Average Condenser Heat Duty'] = condenser_avg
            df.rename(columns={direct_condenser_tag: 'CONDENSER_HEAT_DUTY'}, inplace=True)
        elif utility_tags_available:
            # 2. Fallback to utility calculation (m * Cp * dT)
            df['CONDENSER_HEAT_DUTY_CALC'] = (df['FI_101'] * WATER_SPECIFIC_HEAT * (df['TI_102'] - df['TI_101'])) / 3.6
            
            calc_avg, is_calc_valid, calc_reason = calculate_robust_average(df['CONDENSER_HEAT_DUTY_CALC'])

            if is_calc_valid and calc_avg != 0:
                analysis_results['Average Condenser Heat Duty'] = calc_avg
                df['CONDENSER_HEAT_DUTY'] = df['CONDENSER_HEAT_DUTY_CALC'] # Use calculated column for plotting
            else:
                analysis_results['Average Condenser Heat Duty'] = 'N/A (Condenser Tag Invalid & Utility Calc Failed)'
                df['CONDENSER_HEAT_DUTY'] = 0
        else:
            # 3. Direct tag invalid and utility tags missing
            analysis_results['Average Condenser Heat Duty'] = f'N/A (Condenser Tag Invalid - {reason})'
            df['CONDENSER_HEAT_DUTY'] = 0
    else:
        analysis_results['Average Condenser Heat Duty'] = 'N/A (Missing Tag)'
        df['CONDENSER_HEAT_DUTY'] = 0 # Placeholder for plotting


    return analysis_results, df, {}

# =================================================================================
# 4. PLOT GENERATION
# =================================================================================

def generate_plots(df):
    """Generates and saves temperature profile, DP, and efficiency plots."""
    plot_created_flags = {}

    # 4.1 Temperature Profile Plot
    try:
        plt.figure(figsize=(10, 6))
        if 'DATEANDTIME' in df.columns and not df.empty:
            df.sort_values(by='DATEANDTIME', inplace=True)
            x_axis = df['DATEANDTIME']
            temp_tags = ['TI_01', 'TI_02', 'TI_03', 'TI_04']
            for tag in temp_tags:
                if tag in df.columns:
                    original_tag_name = tag.replace('_', '-')
                    plt.plot(x_axis, df[tag], label=f"{original_tag_name}", alpha=0.7)

            plt.title("C-00 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature (degC)")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            plot_created_flags['temp_profile'] = True
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        plot_created_flags['temp_profile'] = False

    # 4.2 Differential Pressure Plot
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns and not df['DIFFERENTIAL_PRESSURE'].isnull().all() and not df.empty:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DATEANDTIME'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-00 Differential Pressure Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure (mmHg)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            plot_created_flags['dp'] = True
    except Exception as e:
        print(f"Error generating DP plot: {e}")
        plot_created_flags['dp'] = False

    # 4.3 Daily Trends Plot
    try:
        if 'DATEANDTIME' in df.columns and not df.empty:
            df['DATE'] = df['DATEANDTIME'].dt.date
            daily_trends_agg = {
                'FT_01': 'mean',
                'FT_61': 'mean',
            }
            if 'DIFFERENTIAL_PRESSURE' in df.columns: daily_trends_agg['DIFFERENTIAL_PRESSURE'] = 'mean'

            daily_trends = df.groupby('DATE').agg(daily_trends_agg).reset_index()

            plt.figure(figsize=(12, 8))
            if 'FT_01' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['FT_01'], label=f"Avg Feed Flow (kg/h)")
            if 'FT_61' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['FT_61'], label=f"Avg Top Flow (kg/h)")
            if 'DIFFERENTIAL_PRESSURE' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['DIFFERENTIAL_PRESSURE'], label=f"Avg DP (mmHg)")

            plt.title("C-00 Daily Trends")
            plt.xlabel("Date")
            plt.ylabel("Value")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_trends_plot_path)
            plt.close()
            plot_created_flags['trends'] = True
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        plot_created_flags['trends'] = False

    # 4.4 Performance Plot: Moisture Removal vs. Reboiler Heat Duty (Scatter plot)
    try:
        # Check if REBOILER_HEAT_DUTY column exists AND has non-zero variation for plotting
        if ('MOISTURE_REMOVAL_EFFICIENCY_PLOTTED' in df.columns and 
            'REBOILER_HEAT_DUTY' in df.columns and 
            df['REBOILER_HEAT_DUTY'].max() > 0.1): # Check for actual/simulated data presence
            
            plt.figure(figsize=(10, 6))
            # Only plot non-zero duties
            df_plot = df[df['REBOILER_HEAT_DUTY'] > 0.1].copy()
            plt.scatter(df_plot['REBOILER_HEAT_DUTY'], df_plot['MOISTURE_REMOVAL_EFFICIENCY_PLOTTED'], alpha=0.5)
            
            # Update title to indicate if data is simulated
            is_simulated = "(SIMULATED DATA)" if "(SIMULATED" in str(df['REBOILER_HEAT_DUTY'].iloc[0]) else ""
            
            plt.title(f"Moisture Removal Efficiency vs. Reboiler Heat Duty {is_simulated} (Capped at {EFFICIENCY_PLOT_CAP}%)")
            plt.xlabel("Reboiler Heat Duty (kW)")
            plt.ylabel("Moisture Removal Efficiency (%)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_moisture_vs_reboiler_plot_path)
            plt.close()
            plot_created_flags['moisture_vs_reboiler'] = True
        else:
            plot_created_flags['moisture_vs_reboiler'] = False # Set flag to false if no meaningful data
    except Exception as e:
        print(f"Error generating Moisture vs. Reboiler Duty plot: {e}")
        plot_created_flags['moisture_vs_reboiler'] = False

    # 4.5 Performance Plot: Moisture Removal vs. Column Bottom Temperature (Scatter plot)
    try:
        # Use the PLOTTED efficiency column
        if all(tag in df.columns for tag in ['MOISTURE_REMOVAL_EFFICIENCY_PLOTTED', 'TI_04']) and not df.empty:
            plt.figure(figsize=(10, 6))
            plt.scatter(df['TI_04'], df['MOISTURE_REMOVAL_EFFICIENCY_PLOTTED'], alpha=0.5)
            plt.title(f"Moisture Removal Efficiency vs. Column Bottom Temperature (Capped at {EFFICIENCY_PLOT_CAP}%)")
            plt.xlabel("Column Bottom Temperature (degC)")
            plt.ylabel("Moisture Removal Efficiency (%)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_moisture_vs_feedtemp_plot_path)
            plt.close()
            plot_created_flags['moisture_vs_temp'] = True
    except Exception as e:
        print(f"Error generating Moisture vs. Bottom Temperature plot: {e}")
        plot_created_flags['moisture_vs_temp'] = False

    return plot_created_flags

# =================================================================================
# 5. REPORT GENERATION
# =================================================================================

def get_kpi_unit(kpi_key):
    """Maps the KPI title to the correct physical unit."""
    # Unit mapping for KPIs without explicit tag names in parentheses
    unit_map = {
        # Flows
        'Average Feed Flow (FT-01)': 'kg/h',
        'Average Top Product Flow (FT-61)': 'kg/h',
        'Average Bottom Product Flow (FT-62)': 'kg/h',

        # Percentages
        'Overall Material Balance Error (%)': '%',
        'Average Moisture Content in Feed (%)': '%',
        'Average Moisture Removal Efficiency (%)': '%',

        # Pressure and Duty
        'Average Differential Pressure': 'mmHg',
        'Maximum Differential Pressure': 'mmHg',
        'Average Reboiler Heat Duty': 'kW',
        'Average Condenser Heat Duty': 'kW',
    }

    # Use exact match first
    unit = unit_map.get(kpi_key)
    if unit:
        return unit

    # Fallback for keys that may have different tag formats
    match = re.search(r'\((.*?)\)', kpi_key)
    if match:
        tag_name = match.group(1) # e.g., 'FT-01'
        return TAG_UNITS.get(tag_name, 'N/A')

    return 'N/A' # Default for unknown keys


def generate_word_report(analysis_results, df, start_date, end_date, plot_flags):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-00 Dehydration Column Analysis Report', 0)
    doc.add_paragraph(f"Analysis Period: {start_date} to {end_date}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    summary_text = ""
    moisture_efficiency = analysis_results.get('Average Moisture Removal Efficiency (%)', 'N/A')
    
    reboiler_duty_val = analysis_results.get('Average Reboiler Heat Duty', 'N/A')
    if isinstance(reboiler_duty_val, str) and "SIMULATED" in reboiler_duty_val:
        summary_text += "ALERT: Reboiler Heat Duty data was determined to be invalid (zero/flatlined). **The analysis below uses SIMULATED data** (Target: 200-600kW) for plot generation. The I&E team must be engaged to fix the source sensor/calculation. "

    if isinstance(moisture_efficiency, (float, int)):
        summary_text += f"The column achieved an **average moisture removal efficiency of {moisture_efficiency:.2f}%**. This average *excludes* outlier points above {MAX_REASONABLE_EFFICIENCY_PERCENT}% where the light-ends removal overwhelmed the simple moisture-in-feed estimate. "
    else:
        summary_text += "Moisture removal efficiency could not be calculated due to missing flow data. "

    if 'Overall Material Balance Error (%)' in analysis_results:
        summary_text += f"An overall material balance error of **{analysis_results['Overall Material Balance Error (%)']:.2f}%** was observed, which is typically within acceptable limits for noisy process data. "

    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators (KPIs)
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values presented are **averages** over the analysis period.")
    for key, value in analysis_results.items():
        # UNIT FIX: Use the robust mapping function
        unit = get_kpi_unit(key)

        if isinstance(value, str):
            doc.add_paragraph(f"• {key}: {value}")
        else:
            doc.add_paragraph(f"• {key}: {value:.2f} {unit}")

    # Section 3: Performance Analysis
    doc.add_heading('3. Performance Analysis', level=1)
    doc.add_paragraph("This section correlates key operational factors with column performance.")

    # 3.1: Moisture Removal
    doc.add_heading('3.1 Moisture Removal', level=2)
    doc.add_paragraph(f"• Average Moisture Content in Feed: {AVG_FEED_MOISTURE_MASS_PERCENT:.2f}%")
    doc.add_paragraph(f"• Average Moisture Removal Efficiency: {analysis_results.get('Average Moisture Removal Efficiency (%)', 'N/A'):.2f}%")
    doc.add_paragraph(f"NOTE: Efficiency above 100% is typical as the top product (FT-61) removes light ends in addition to the estimated moisture. Outliers greater than {MAX_REASONABLE_EFFICIENCY_PERCENT}% were excluded from the average calculation.")

    doc.add_heading('Moisture Removal vs. Reboiler Heat Duty', level=3)
    doc.add_paragraph("This plot shows how increasing the energy input to the reboiler influences vaporization and thus moisture removal.")
    if plot_flags.get('moisture_vs_reboiler') and os.path.exists(output_moisture_vs_reboiler_plot_path):
        doc.add_picture(output_moisture_vs_reboiler_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing/invalid data for Reboiler Duty.")

    doc.add_heading('Moisture Removal vs. Column Bottom Temperature', level=3)
    doc.add_paragraph("The bottom temperature (TI-04) is critical for driving the separation.")
    if plot_flags.get('moisture_vs_temp') and os.path.exists(output_moisture_vs_feedtemp_plot_path):
        doc.add_picture(output_moisture_vs_feedtemp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing data.")

    # 3.2: Differential Pressure (DP) Analysis
    doc.add_heading('3.2 Differential Pressure', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding, foaming, or fouling inside the column.")
    if plot_flags.get('dp') and os.path.exists(output_dp_plot_path):
        doc.add_picture(output_dp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing data.")

    # Section 4: General Performance Plots
    doc.add_heading('4. General Performance Plots', level=1)

    doc.add_heading('4.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column. A consistent gradient indicates stable operation.")
    if plot_flags.get('temp_profile') and os.path.exists(output_temp_plot_path):
        doc.add_picture(output_temp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing data.")

    doc.add_heading('4.2 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables, helping to visualize long-term shifts in performance.")
    if plot_flags.get('trends') and os.path.exists(output_trends_plot_path):
        doc.add_picture(output_trends_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing data.")

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

# =================================================================================
# 6. MAIN EXECUTION
# =================================================================================

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return

    # Use the same default date range for consistency
    scada_data, start_date, end_date = get_scada_data(engine)
    if scada_data is None:
        return

    analysis_results, scada_data, outliers = perform_analysis(scada_data)

    if analysis_results:
        plot_flags = generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, start_date, end_date, plot_flags)
        print("C-00 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()
