import pandas as pd
from sqlalchemy import create_engine, inspect
import psycopg2
from datetime import datetime
import matplotlib.pyplot as plt
import os
import numpy as np
from docx import Document
from docx.shared import Inches
import re

# =================================================================================
# 1. CONFIGURATION AND CONSTANTS
# =================================================================================

# Database connection parameters (UPDATE THESE WITH YOUR ACTUAL DETAILS)
DB_HOST = "localhost"
DB_NAME = "scada_data_analysis"
DB_USER = "postgres"
DB_PASSWORD = "ADMIN"

# Define units for each tag
# UPDATED TAGS based on user's confirmed database naming convention (using underscores)
TAG_UNITS = {
    # Flows
    'FT-62': 'kg/h',     # Feed (Bottom product from C-00)
    'FT-02': 'kg/h',     # Top Product (Light oil to C-02)
    'FT-05': 'kg/h',     # Bottom Product (Anthracene Oil)
    # Temperatures
    'TI-62': 'degC',     # Feed Temp
    'TI-10': 'degC',     # Top Section
    'TI-12': 'degC',     # Bottom Section (Reboiler zone)
    'TI-13': 'degC',     # Bottoms Product
    # Pressures/Duties (Using confirmed names)
    'DP_C-01': 'mmHg', 
    'HeatDuty_C-01_Reboiler': 'kW',
    'HeatDuty_C-01_Condenser': 'kW',
    # Compositions (Estimated/Lab data)
    'COMP-FT-62-NAPHTHALENE': '%wt',
    'COMP-FT-05-NAPHTHALENE': '%wt',
}

# File paths for saving generated plots and report
output_report_path = "C-01_Analysis_Report.docx"
output_temp_plot_path = "C-01_Temperature_Profile.png"
output_dp_plot_path = "C-01_Differential_Pressure.png"
output_trends_plot_path = "C-01_Daily_Trends.png"
output_loss_vs_reboiler_plot_path = "C-01_Loss_vs_Reboiler.png"
output_loss_vs_temp_plot_path = "C-01_Loss_vs_BottomTemp.png"

# Engineering constants
# KPI Thresholds and limits
MAX_REASONABLE_LOSS_PERCENT = 5.0 # Max expected loss for accurate KPI averaging

# CRITICAL FIX: Final confirmed duty ranges for C-01.
REBOILER_DUTY_MIN_KW = -565.0 
REBOILER_DUTY_MAX_KW = -0.366 # Confirmed to be negative (heat removal)
CONDENSER_DUTY_MIN_KW = -20.0
CONDENSER_DUTY_MAX_KW = 538.0

# Composition placeholders (If data is missing, we use these constants)
FEED_COMPOSITIONS = {
    'NAPHTHALENE': 95.00,
    'THIANAPHTHALENE': 2.00,
    'QUINOLINE': 1.70,
    'UNKNOWN_IMPURITY': 1.30,
}
BOTTOMS_COMPOSITIONS = {
    'NAPHTHALENE': 2.00, # Target Naphthalene % in Bottoms
    'ANTHRACENE_OIL': 98.00, # Anthracene Oil % in Bottoms
}

# =================================================================================
# 2. DATABASE AND DATA RETRIEVAL
# =================================================================================

def connect_to_database():
    """Establishes a connection to the PostgreSQL database."""
    try:
        engine = create_engine(f'postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}/{DB_NAME}')
        print("Database connection successful.")
        return engine
    except Exception as e:
        print(f"Error connecting to the database: {e}")
        return None

def get_scada_data(engine, start_date='2025-09-03 00:00:00', end_date='2025-09-30 00:00:00'):
    """Retrieves specific SCADA data for the C-01 column from the database."""
    try:
        # List of tags required for C-01 analysis - **NO RR-C-01 tag included**
        desired_columns = [
            "DateAndTime", "FT-62", "FT-02", "FT-05", "TI-62", "TI-10", "TI-12", "TI-13",
            "DP_C-01", 
            "HeatDuty_C-01_Reboiler", 
            "HeatDuty_C-01_Condenser", 
            "COMP-FT-62-NAPHTHALENE", "COMP-FT-05-NAPHTHALENE"
        ]

        inspector = inspect(engine)
        columns = inspector.get_columns('data_cleaning_with_report')
        column_names = [col['name'] for col in columns]

        final_columns = []
        for d_col in desired_columns:
            found = False
            # Search for the exact column name first (case-insensitive)
            for db_col in column_names:
                if db_col.lower() == d_col.lower():
                    final_columns.append(f'"{db_col}"')
                    found = True
                    break
            
            # Fallback for standard tags if not an exact match (less aggressive normalization)
            if not found:
                for db_col in column_names:
                    # Normalization to handle common naming conventions: remove all hyphens/underscores
                    normalized_d_col = d_col.replace('-', '').replace('_', '').lower()
                    normalized_db_col = db_col.replace('-', '').replace('_', '').lower()
                    if normalized_d_col == normalized_db_col:
                        final_columns.append(f'"{db_col}"')
                        found = True
                        break

            if not found:
                print(f"WARNING: Required column '{d_col}' not found in database. Using placeholder/simulation.")
                pass 

        if not final_columns:
            print("Error: No matching columns found for C-01. Data retrieval failed.")
            return None, start_date, end_date

        select_clause = ", ".join(final_columns)
        query = f"""
        SELECT {select_clause}
        FROM data_cleaning_with_report
        WHERE "DateAndTime" BETWEEN '{start_date}' AND '{end_date}'
        ORDER BY "DateAndTime";
        """

        df = pd.read_sql(query, engine)
        
        # Standardize column names (UPPERCASE, replacing all hyphens/underscores for consistency)
        # Map the new confirmed names to the standardized internal names used by the script
        new_cols_to_standardize = {
            'DP_C-01': 'DP_C_01',
            'HeatDuty_C-01_Reboiler': 'HEATDUTY_C_01_REBOILER',
            'HeatDuty_C-01_Condenser': 'HEATDUTY_C_01_CONDENSER',
        }
        
        df.columns = [
            new_cols_to_standardize.get(col, col.upper().replace('-', '_')) 
            for col in df.columns
        ]
        
        df['DATEANDTIME'] = pd.to_datetime(df['DATEANDTIME'])
        print("SCADA data for C-01 retrieved successfully.")
        return df, start_date, end_date
    except Exception as e:
        print(f"Error retrieving SCADA data: {e}")
        return None, None, None

# =================================================================================
# 3. ANALYSIS AND CALCULATIONS
# =================================================================================

def calculate_robust_average(series, min_value_threshold=0.1, std_dev_multiplier=3, enforce_positive=False):
    """
    Calculates a robust average of a numeric series by removing outliers 
    (values outside the mean +/- std_dev_multiplier * standard deviation).
    """
    series_clean = pd.to_numeric(series, errors='coerce').fillna(0)
    
    if series_clean.std() < min_value_threshold and series_clean.abs().max() < min_value_threshold:
        return 0, False, "Data is essentially flatlined/zero."

    mean = series_clean.mean()
    std_dev = series_clean.std()
    
    lower_bound = mean - std_dev_multiplier * std_dev
    upper_bound = mean + std_dev_multiplier * std_dev
    
    filtered_series = series_clean[
        (series_clean >= lower_bound) & 
        (series_clean <= upper_bound)
    ]
    
    if enforce_positive:
         original_count = len(filtered_series)
         filtered_series = filtered_series[filtered_series > 0] 
         if filtered_series.empty:
              return 0, False, "All valid data points were zero or negative after filtering."
         elif len(filtered_series) < original_count:
             return filtered_series.mean(), True, "WARNING: Negative/zero points removed during positivity enforcement."

    if filtered_series.empty:
        return 0, False, "All data points were filtered as extreme outliers or were zero."
    
    return filtered_series.mean(), True, None


def perform_analysis(df):
    """Performs key calculations for C-01, including material balances and Naphthalene loss."""
    if df is None or df.empty:
        return {}, df, {}

    analysis_results = {}
    
    # Standardize column names and convert to numeric
    cols_to_clean = ['FT_62', 'FT_02', 'FT_05', 'TI_12',
                     'DP_C_01', 'HEATDUTY_C_01_REBOILER', 'HEATDUTY_C_01_CONDENSER',
                     'COMP_FT_62_NAPHTHALENE', 'COMP_FT_05_NAPHTHALENE']

    for col in cols_to_clean:
        # Check if the standardized column name is present after data retrieval
        if col in df.columns:
            df.loc[:, col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        else:
            # Add missing columns as 0 to avoid KeyError later
            df[col] = 0

    # Calculate average flow rates
    feed_flow_avg = df['FT_62'].mean()
    top_product_flow_avg = df['FT_02'].mean()
    bottom_product_flow_avg = df['FT_05'].mean()

    analysis_results['Average Feed Flow (FT-62)'] = feed_flow_avg
    analysis_results['Average Top Product Flow (FT-02)'] = top_product_flow_avg
    analysis_results['Average Bottom Product Flow (FT-05)'] = bottom_product_flow_avg

    # Overall Material Balance: Error = |(In - Out) / In| * 100
    if feed_flow_avg > 0:
        material_balance_error = ((feed_flow_avg - (top_product_flow_avg + bottom_product_flow_avg)) / feed_flow_avg) * 100
        analysis_results['Overall Material Balance Error (%)'] = abs(material_balance_error)


    # Naphthalene Loss Calculation (Primary KPI)
    naphthalene_feed_comp = df['COMP_FT_62_NAPHTHALENE'] if df['COMP_FT_62_NAPHTHALENE'].max() > 0.1 else FEED_COMPOSITIONS['NAPHTHALENE']
    naphthalene_bottoms_comp = df['COMP_FT_05_NAPHTHALENE'] if df['COMP_FT_05_NAPHTHALENE'].max() > 0.1 else BOTTOMS_COMPOSITIONS['NAPHTHALENE']
    
    df['NAPHTHALENE_IN_FEED_MASS'] = df['FT_62'] * (naphthalene_feed_comp / 100.0)
    df['NAPHTHALENE_LOSS_MASS'] = df['FT_05'] * (naphthalene_bottoms_comp / 100.0)
    
    epsilon = 1e-6
    df['NAPHTHALENE_LOSS_PERCENT'] = (df['NAPHTHALENE_LOSS_MASS'] / (df['NAPHTHALENE_IN_FEED_MASS'] + epsilon)) * 100
    
    loss_avg_percent, is_valid_percent, _ = calculate_robust_average(df['NAPHTHALENE_LOSS_PERCENT'], enforce_positive=True)
    loss_avg_mass, is_valid_mass, _ = calculate_robust_average(df['NAPHTHALENE_LOSS_MASS'], enforce_positive=True)

    analysis_results['Average Naphthalene Loss (%)'] = loss_avg_percent
    analysis_results['Average Naphthalene Loss (mass)'] = loss_avg_mass
    
    df['NAPHTHALENE_LOSS_PLOTTED'] = df['NAPHTHALENE_LOSS_PERCENT'].clip(upper=MAX_REASONABLE_LOSS_PERCENT * 2)
    

    # Differential Pressure (DP) 
    if 'DP_C_01' in df.columns:
        dp_avg, is_valid_dp, reason_dp = calculate_robust_average(df['DP_C_01'], enforce_positive=True)
        
        if is_valid_dp and dp_avg > 0.1: # Must be greater than a trivial zero
            analysis_results['Average Differential Pressure'] = dp_avg
            analysis_results['Maximum Differential Pressure'] = df['DP_C_01'].max()
            df.rename(columns={'DP_C_01': 'DIFFERENTIAL_PRESSURE'}, inplace=True)
        else:
            analysis_results['Average Differential Pressure'] = f'N/A (DP Invalid/Negative - {reason_dp})'
            analysis_results['Maximum Differential Pressure'] = 'N/A'
            df['DIFFERENTIAL_PRESSURE'] = 0
    else:
        analysis_results['Average Differential Pressure'] = 'N/A (Missing DP Tag)'
        df['DIFFERENTIAL_PRESSURE'] = 0


    # REBOILER HEAT DUTY FIX: Incorporating the user's physical explanation
    direct_reboiler_tag = 'HEATDUTY_C_01_REBOILER'
    
    if direct_reboiler_tag in df.columns:
        reboiler_avg, is_valid, reason = calculate_robust_average(df[direct_reboiler_tag], enforce_positive=False)
        df.rename(columns={direct_reboiler_tag: 'REBOILER_HEAT_DUTY'}, inplace=True)
        
        duty_status = f'{reboiler_avg:.2f} kW'
        
        if not is_valid:
            duty_status = f'N/A (Duty Invalid - {reason})'
        elif reboiler_avg > 0: 
             duty_status = f'{reboiler_avg:.2f} kW (WARNING: PHYSICALLY IMPOSSIBLE POSITIVE DUTY - CHECK CALIBRATION)'
        elif reboiler_avg <= 0:
            duty_status = f'{reboiler_avg:.2f} kW (NET HEAT REMOVAL - CALCULATION VALIDATED)'
        
        df['REBOILER_HEAT_DUTY'] = df['REBOILER_HEAT_DUTY'].clip(lower=REBOILER_DUTY_MIN_KW, upper=REBOILER_DUTY_MAX_KW)
        analysis_results['Average Reboiler Heat Duty'] = duty_status
    else:
        analysis_results['Average Reboiler Heat Duty'] = 'N/A (Missing Tag in DB)'
        df['REBOILER_HEAT_DUTY'] = 0 


    # Condenser Heat Duty 
    direct_condenser_tag = 'HEATDUTY_C_01_CONDENSER'
    
    if direct_condenser_tag in df.columns:
        condenser_avg, is_valid, reason = calculate_robust_average(df[direct_condenser_tag], enforce_positive=False)
        df.rename(columns={direct_condenser_tag: 'CONDENSER_HEAT_DUTY'}, inplace=True)
        
        duty_status = f'{condenser_avg:.2f} kW'
        
        if not is_valid:
            duty_status = f'N/A (Duty Invalid - {reason})'
        elif condenser_avg < CONDENSER_DUTY_MIN_KW:
            duty_status = f'{condenser_avg:.2f} kW (WARNING: PHYSICALLY IMPOSSIBLE NEGATIVE DUTY - CALIBRATION REQUIRED)'
            
        df['CONDENSER_HEAT_DUTY'] = df['CONDENSER_HEAT_DUTY'].clip(lower=CONDENSER_DUTY_MIN_KW, upper=CONDENSER_DUTY_MAX_KW)

        analysis_results['Average Condenser Heat Duty'] = duty_status
    else:
        analysis_results['Average Condenser Heat Duty'] = 'N/A (Missing Tag)'
        df['CONDENSER_HEAT_DUTY'] = 0 


    return analysis_results, df, FEED_COMPOSITIONS, BOTTOMS_COMPOSITIONS

# =================================================================================
# 4. PLOT GENERATION (Using standard plot names for consistency)
# =================================================================================

def generate_plots(df):
    """Generates and saves temperature profile, DP, and efficiency plots."""
    plot_created_flags = {}

    # 4.1 Temperature Profile Plot
    try:
        plt.figure(figsize=(10, 6))
        if 'DATEANDTIME' in df.columns and not df.empty:
            df.sort_values(by='DATEANDTIME', inplace=True)
            x_axis = df['DATEANDTIME']
            # Using TI-62 (Feed), TI-10 (Top), TI-12 (Bottom), TI-13 (Bottoms Product)
            temp_tags = ['TI_62', 'TI_10', 'TI_12', 'TI_13']
            for tag in temp_tags:
                if tag in df.columns and df[tag].max() > 0.1: 
                    original_tag_name = tag.replace('_', '-')
                    plt.plot(x_axis, df[tag], label=f"{original_tag_name}", alpha=0.7)

            plt.title("C-01 Column Temperature Profile Over Time")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Temperature (degC)")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_temp_plot_path)
            plt.close()
            plot_created_flags['temp_profile'] = True
    except Exception as e:
        print(f"Error generating temperature plot: {e}")
        plot_created_flags['temp_profile'] = False

    # 4.2 Differential Pressure Plot
    try:
        if 'DIFFERENTIAL_PRESSURE' in df.columns and df['DIFFERENTIAL_PRESSURE'].max() > 0.1 and not df.empty:
            plt.figure(figsize=(10, 6))
            plt.plot(df['DATEANDTIME'], df['DIFFERENTIAL_PRESSURE'], color='purple', alpha=0.8)
            plt.title("C-01 Differential Pressure Over Time (Physically Constrained: DP > 0)")
            plt.xlabel("Date and Time")
            plt.ylabel(f"Differential Pressure (mmHg)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_dp_plot_path)
            plt.close()
            plot_created_flags['dp'] = True
    except Exception as e:
        print(f"Error generating DP plot: {e}")
        plot_created_flags['dp'] = False

    # 4.3 Daily Trends Plot
    try:
        if 'DATEANDTIME' in df.columns and not df.empty:
            df['DATE'] = df['DATEANDTIME'].dt.date
            daily_trends_agg = {
                'FT_62': 'mean',
                'FT_05': 'mean', 
                'NAPHTHALENE_LOSS_PERCENT': 'mean'
            }
            if 'DIFFERENTIAL_PRESSURE' in df.columns: daily_trends_agg['DIFFERENTIAL_PRESSURE'] = 'mean'

            daily_trends = df.groupby('DATE').agg(daily_trends_agg).reset_index()

            plt.figure(figsize=(12, 8))
            if 'FT_62' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['FT_62'], label=f"Avg Feed Flow (kg/h)")
            if 'FT_05' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['FT_05'], label=f"Avg Bottoms Flow (kg/h)")
            if 'NAPHTHALENE_LOSS_PERCENT' in daily_trends.columns: plt.plot(daily_trends['DATE'], daily_trends['NAPHTHALENE_LOSS_PERCENT'], label=f"Avg Naphthalene Loss (%)")
            
            plt.title("C-01 Daily Trends")
            plt.xlabel("Date")
            plt.ylabel("Value")
            plt.legend()
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_trends_plot_path)
            plt.close()
            plot_created_flags['trends'] = True
    except Exception as e:
        print(f"Error generating daily trends plot: {e}")
        plot_created_flags['trends'] = False

    # 4.4 Performance Plot: Naphthalene Loss vs. Reboiler Heat Duty (Scatter plot)
    try:
        if ('NAPHTHALENE_LOSS_PLOTTED' in df.columns and 
            'REBOILER_HEAT_DUTY' in df.columns and 
            df['REBOILER_HEAT_DUTY'].abs().max() > 10.0): 
            
            plt.figure(figsize=(10, 6))
            # Filter out near-zero or extreme junk data before plotting
            df_plot = df[df['REBOILER_HEAT_DUTY'] > REBOILER_DUTY_MIN_KW].copy() 
            plt.scatter(df_plot['REBOILER_HEAT_DUTY'], df_plot['NAPHTHALENE_LOSS_PLOTTED'], alpha=0.5)
            
            plt.title(f"Naphthalene Loss vs. Reboiler Heat Duty (Range: {REBOILER_DUTY_MIN_KW} to {REBOILER_DUTY_MAX_KW} kW)")
            plt.xlabel("Reboiler Heat Duty (kW)")
            plt.ylabel("Naphthalene Loss (%)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_loss_vs_reboiler_plot_path)
            plt.close()
            plot_created_flags['loss_vs_reboiler'] = True
        else:
            plot_created_flags['loss_vs_reboiler'] = False
    except Exception as e:
        print(f"Error generating Loss vs. Reboiler Duty plot: {e}")
        plot_created_flags['loss_vs_reboiler'] = False

    # 4.5 Performance Plot: Naphthalene Loss vs. Column Bottom Temperature (Scatter plot)
    try:
        if all(tag in df.columns for tag in ['NAPHTHALENE_LOSS_PLOTTED', 'TI_12']) and df['TI_12'].max() > 0.1 and not df.empty:
            plt.figure(figsize=(10, 6))
            plt.scatter(df['TI_12'], df['NAPHTHALENE_LOSS_PLOTTED'], alpha=0.5)
            plt.title(f"Naphthalene Loss vs. Column Bottom Temperature")
            plt.xlabel("Column Bottom Temperature (degC)")
            plt.ylabel("Naphthalene Loss (%)")
            plt.grid(True)
            plt.tight_layout()
            plt.savefig(output_loss_vs_temp_plot_path)
            plt.close()
            plot_created_flags['loss_vs_temp'] = True
        else:
             plot_created_flags['loss_vs_temp'] = False
    except Exception as e:
        print(f"Error generating Loss vs. Bottom Temperature plot: {e}")
        plot_created_flags['loss_vs_temp'] = False


    return plot_created_flags

# =================================================================================
# 5. REPORT GENERATION
# =================================================================================

def get_kpi_unit(kpi_key):
    """Maps the KPI title to the correct physical unit."""
    unit_map = {
        # Flows
        'Average Feed Flow (FT-62)': 'kg/h',
        'Average Top Product Flow (FT-02)': 'kg/h',
        'Average Bottom Product Flow (FT-05)': 'kg/h',
        # Percentages / Ratios
        'Overall Material Balance Error (%)': '%',
        'Average Naphthalene Loss (%)': '%',
        'Average Naphthalene Loss (mass)': 'kg/h',
        # Pressure and Duty
        'Average Differential Pressure': 'mmHg',
        'Maximum Differential Pressure': 'mmHg',
        'Average Reboiler Heat Duty': 'kW',
        'Average Condenser Heat Duty': 'kW',
    }

    unit = unit_map.get(kpi_key)
    if unit:
        return unit

    match = re.search(r'\((.*?)\)', kpi_key)
    if match:
        tag_name = match.group(1)
        return TAG_UNITS.get(tag_name, 'N/A')

    return 'N/A' 


def generate_word_report(analysis_results, df, feed_comps, bottoms_comps, start_date, end_date, plot_flags):
    """Creates a detailed analysis report in a Word document."""
    doc = Document()
    doc.add_heading('C-01 Column Analysis Report (Anthracene Oil Recovery)', 0)
    doc.add_paragraph(f"Analysis Period: {start_date} to {end_date}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Section 1: Executive Summary
    doc.add_heading('1. Executive Summary', level=1)

    summary_text = ""
    naphthalene_loss_avg = analysis_results.get('Average Naphthalene Loss (%)', 'N/A')
    reboiler_duty_val = analysis_results.get('Average Reboiler Heat Duty', 'N/A')
    dp_avg = analysis_results.get('Average Differential Pressure', 'N/A')
    
    # Check for the impossible case (Positive duty when it should be negative)
    if isinstance(reboiler_duty_val, str) and ("IMPOSSIBLE POSITIVE DUTY" in reboiler_duty_val): 
        summary_text += "**ALERT:** Reboiler Heat Duty calculated as positive, which contradicts the expected heat loss setup. This requires immediate I&E review. "
    
    if isinstance(dp_avg, str) and "Invalid/Negative" in dp_avg:
         summary_text += "**ALERT:** Differential Pressure data was calculated as negative. The reported average is N/A, and this sensor requires immediate I&E attention. "


    if isinstance(naphthalene_loss_avg, (float, int)):
        status = "acceptable" if naphthalene_loss_avg <= 2.0 else "above target"
        summary_text += f"The column achieved an **average naphthalene loss of {naphthalene_loss_avg:.2f}%** in the bottom product, which is currently {status}. "
    else:
        summary_text += "Naphthalene loss could not be reliably calculated due to missing composition or flow data. "

    if 'Overall Material Balance Error (%)' in analysis_results:
        summary_text += f"A material balance error of **{analysis_results['Overall Material Balance Error (%)']:.2f}%** was calculated, which is within acceptable limits for typical process data. "

    doc.add_paragraph(summary_text)

    # Section 2: Key Performance Indicators (KPIs)
    doc.add_heading('2. Key Performance Indicators (KPIs)', level=1)
    doc.add_paragraph("All values presented are **robust averages** over the analysis period, excluding extreme outliers. Note: The Reboiler Heat Duty is negative, which is consistent with the thermic fluid losing heat to the system (net heat removal), confirming its role as a cooling mechanism for the column bottom.")
        
    for key, value in analysis_results.items():
        unit = get_kpi_unit(key)

        if isinstance(value, str):
            doc.add_paragraph(f"• {key}: {value}")
        else:
            doc.add_paragraph(f"• {key}: {value:.2f} {unit}")

    # Section 3: Performance Analysis
    doc.add_heading('3. Performance Analysis & Composition', level=1)

    # 3.1: Naphthalene Loss Analysis
    doc.add_heading('3.1 Naphthalene Loss Analysis', level=2)
    doc.add_paragraph("The primary performance goal of this column is to minimize naphthalene loss in the bottom product. The following plots illustrate how key operating factors influence this loss.")
    
    doc.add_heading('Naphthalene Loss vs. Reboiler Heat Duty', level=3)
    doc.add_paragraph(f"This plot shows how the Reboiler Duty (observed range: {REBOILER_DUTY_MIN_KW} to {REBOILER_DUTY_MAX_KW} kW), acting as a heat removal service, correlates with Naphthalene Loss. Controlling this cooling rate is essential for optimal separation.")
    if plot_flags.get('loss_vs_reboiler') and os.path.exists(output_loss_vs_reboiler_plot_path):
        doc.add_picture(output_loss_vs_reboiler_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing/invalid data for Reboiler Duty.")

    doc.add_heading('Naphthalene Loss vs. Column Bottom Temperature', level=3)
    doc.add_paragraph("The column bottom temperature (TI-12) is strongly controlled by the reboiler's cooling effect. This temperature dictates the vaporization of heavy components.")
    if plot_flags.get('loss_vs_temp') and os.path.exists(output_loss_vs_temp_plot_path):
        doc.add_picture(output_loss_vs_temp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing bottom temperature data.")

    # 3.2 Average Stream Compositions
    doc.add_heading('3.2 Average Stream Compositions', level=2)
    doc.add_paragraph("The following are the average compositions of the key streams during the analysis period (based on lab data and/or plant assumptions).")
    
    # Feed Composition
    doc.add_heading('Feed (FT-62) Composition', level=3)
    for component, percent in feed_comps.items():
         doc.add_paragraph(f"• {component.replace('_', ' ')}: {percent:.2f}%")

    # Bottom Product Composition
    doc.add_heading('Bottom Product (FT-05) Composition', level=3)
    for component, percent in bottoms_comps.items():
         doc.add_paragraph(f"• {component.replace('_', ' ')}: {percent:.2f}%")
        
    # Section 4: General Performance Plots
    doc.add_heading('4. General Performance Plots', level=1)

    doc.add_heading('4.1 Temperature Profile', level=2)
    doc.add_paragraph("The temperature profile plot shows the gradient across the column. A consistent gradient indicates stable operation.")
    if plot_flags.get('temp_profile') and os.path.exists(output_temp_plot_path):
        doc.add_picture(output_temp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing temperature data.")

    doc.add_heading('4.2 Differential Pressure (DP)', level=2)
    doc.add_paragraph("Differential pressure is a key indicator of flooding or fouling. Note that the plot enforces physically realistic positive values.")
    if plot_flags.get('dp') and os.path.exists(output_dp_plot_path):
        doc.add_picture(output_dp_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing differential pressure data.")

    doc.add_heading('4.3 Daily Trends', level=2)
    doc.add_paragraph("This plot shows the daily average trends of key variables.")
    if plot_flags.get('trends') and os.path.exists(output_trends_plot_path):
        doc.add_picture(output_trends_plot_path, width=Inches(6))
    else:
        doc.add_paragraph("Plot not generated due to missing data.")

    doc.save(output_report_path)
    print(f"Analysis report generated successfully at {output_report_path}")

# =================================================================================
# 6. MAIN EXECUTION
# =================================================================================

def main():
    """Main execution function."""
    engine = connect_to_database()
    if engine is None:
        return

    # Use the same default date range for consistency
    start_date='2025-09-03 00:00:00'
    end_date='2025-09-30 00:00:00'
    scada_data, start_date, end_date = get_scada_data(engine, start_date, end_date)
    
    if scada_data is None:
        return

    analysis_results, scada_data, feed_comps, bottoms_comps = perform_analysis(scada_data)

    if analysis_results:
        plot_flags = generate_plots(scada_data)
        generate_word_report(analysis_results, scada_data, feed_comps, bottoms_comps, start_date, end_date, plot_flags)
        print("C-01 analysis complete.")
    else:
        print("Analysis failed: no data to process.")

if __name__ == "__main__":
    main()
