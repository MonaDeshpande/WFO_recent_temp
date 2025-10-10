import numpy as np
import pandas as pd
from sklearn.ensemble import IsolationForest, RandomForestRegressor
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import os
import random
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches

# Import the main data retrieval function and configuration
from data_retrieval_R00 import main_data_retrieval
from config_R00 import START_DATE, END_DATE, GC_FILE, SCADA_TAGS_C00, SCADA_TAGS_C01, SCADA_TAGS_C02, SCADA_TAGS_C03

# Define the dictionary of relevant SCADA tags for each column.
SCADA_TAGS_BY_COLUMN = {
    'C-00': SCADA_TAGS_C00,
    'C-01': SCADA_TAGS_C01,
    'C-02': SCADA_TAGS_C02,
    'C-03': SCADA_TAGS_C03
}

def create_word_report(column_id, pm_alert_msg, ad_alert_msg, pm_plot_path, ad_plot_path):
    """
    Generates a Word document summarizing the analysis results.
    """
    document = Document()
    document.add_heading(f'Analysis Report for Column {column_id}', level=1)
    
    # Section for Predictive Maintenance
    document.add_heading('1. Predictive Maintenance Analysis', level=2)
    document.add_paragraph('This analysis predicts the future trend of a key sensor to forecast potential equipment failure.')
    document.add_paragraph(f'Status: {pm_alert_msg}')
    if os.path.exists(pm_plot_path):
        document.add_picture(pm_plot_path, width=Inches(6))
    
    # Section for Anomaly Detection
    document.add_heading('2. Real-time Anomaly Detection Analysis', level=2)
    document.add_paragraph('This analysis identifies unusual sensor readings that deviate from normal operating behavior.')
    document.add_paragraph(f'Status: {ad_alert_msg}')
    if os.path.exists(ad_plot_path):
        document.add_picture(ad_plot_path, width=Inches(6))
    
    # Save the document
    report_filename = f'Analysis_Report_{column_id}.docx'
    document.save(report_filename)
    print(f"Report for {column_id} saved to {report_filename}")

def train_models(scada_df, column_id):
    """
    Trains a predictive maintenance model and an anomaly detection model on historical data.
    """
    print(f"\n--- Training Models for {column_id} ---")
    if scada_df is None or scada_df.empty:
        print("No data available to train models.")
        return None, None, None, None, None, None

    # --- Predictive Maintenance Model (Random Forest) ---
    df_pred_maint = scada_df.copy()
    df_pred_maint['Day'] = (df_pred_maint['DATEANDTIME'] - df_pred_maint['DATEANDTIME'].min()).dt.days + 1
    np.random.seed(42)
    base_pressure = 50.0
    noise = np.random.normal(0, 2, len(df_pred_maint))
    df_pred_maint['Pressure_Reading'] = base_pressure + df_pred_maint['Day']**1.5 * 0.05 + noise

    model_pm = RandomForestRegressor(n_estimators=100, random_state=42)
    model_pm.fit(df_pred_maint[['Day']], df_pred_maint['Pressure_Reading'])
    
    pm_plot_path = f"{column_id}_predictive_maintenance_plot.png"
    plt.figure(figsize=(12, 6))
    plt.scatter(df_pred_maint['Day'], df_pred_maint['Pressure_Reading'], label='Historical Data', alpha=0.6)
    plt.plot(df_pred_maint['Day'], model_pm.predict(df_pred_maint[['Day']]), color='red', linestyle='--', label='Random Forest Fit')
    plt.title(f"Predictive Maintenance for {column_id} Sensor Pressure")
    plt.xlabel("Day")
    plt.ylabel("Sensor Pressure (PSI)")
    plt.legend()
    plt.grid(True)
    plt.savefig(pm_plot_path)
    plt.close()

    # --- Anomaly Detection Model (Isolation Forest) ---
    numerical_cols = [col for col in scada_df.columns if col not in ['DATEANDTIME', 'DateAndTime', 'dateandtime']]
    if len(numerical_cols) < 2:
        print(f"Not enough numerical tags for anomaly detection for {column_id}. Skipping.")
        return model_pm, None, None, pm_plot_path, None, numerical_cols

    features = scada_df.dropna(subset=numerical_cols)[numerical_cols].values
    scaler = StandardScaler()
    scaled_features = scaler.fit_transform(features)

    model_ad = IsolationForest(contamination='auto', random_state=42)
    model_ad.fit(scaled_features)
    
    ad_plot_path = f"{column_id}_anomaly_detection_plot.png"
    
    return model_pm, model_ad, scaler, pm_plot_path, ad_plot_path, numerical_cols

def run_real_time_analysis(column_id, scada_df, pm_model, ad_model, scaler, ad_cols):
    """
    Simulates a real-time data stream and performs analysis.
    """
    print(f"\n--- Running Real-time Analysis for {column_id} ---")
    pm_alert_msg = "No immediate failure predicted within the next 50 days."
    ad_alert_msg = "No significant anomalies detected."
    
    critical_threshold = 100.0
    
    pm_failure_found = False
    ad_anomaly_found = False

    # --- Predictive Maintenance Inference (RUN ONCE) ---
    if pm_model:
        last_day = (scada_df['DATEANDTIME'].iloc[-1] - scada_df['DATEANDTIME'].min()).days + 1
        future_days = np.arange(last_day, last_day + 50).reshape(-1, 1)
        predicted_pressure = pm_model.predict(future_days)
        
        if any(p >= critical_threshold for p in predicted_pressure):
            pm_failure_found = True

    # --- Anomaly Detection Inference (RUN IN LOOP) ---
    for i in range(len(scada_df)):
        live_data_point = scada_df.iloc[i:i+1].copy()
        
        if ad_model and ad_cols:
            if live_data_point[ad_cols].notna().any().any():
                live_features = live_data_point[ad_cols].values
                
                # Generate a mock anomaly with the correct number of features
                if i == 50:
                    n_features = len(ad_cols)
                    live_features = np.array([np.random.rand(n_features) * 10]).reshape(1, -1)
                    
                live_scaled = scaler.transform(live_features)
                is_anomaly = ad_model.predict(live_scaled)
                
                if is_anomaly[0] == -1:
                    ad_anomaly_found = True

    if pm_failure_found:
        pm_alert_msg = "ALERT: Failure predicted within the next 50 days!"
    if ad_anomaly_found:
        ad_alert_msg = "ALERT: Anomaly detected in the data stream!"

    return pm_alert_msg, ad_alert_msg
    
def main():
    """Main execution function to run advanced analysis on all columns."""
    for column_id, tags in SCADA_TAGS_BY_COLUMN.items():
        print("="*60)
        print(f"Beginning advanced analysis for {column_id}")
        
        # Step 1: Get data and train models
        scada_data, _ = main_data_retrieval(column_id, tags, START_DATE, END_DATE, GC_FILE, [])
        
        if scada_data is None or scada_data.empty:
            print(f"No SCADA data available for {column_id}. Skipping.")
            continue
            
        pm_model, ad_model, scaler, pm_plot_path, ad_plot_path, ad_cols = train_models(scada_data, column_id)
        
        # Step 2: Run the real-time simulation and get alert messages
        pm_alert_msg, ad_alert_msg = run_real_time_analysis(column_id, scada_data, pm_model, ad_model, scaler, ad_cols)
        
        # Step 3: Generate the final Word report
        create_word_report(column_id, pm_alert_msg, ad_alert_msg, pm_plot_path, ad_plot_path)

if __name__ == "__main__":
    main()