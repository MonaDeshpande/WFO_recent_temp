import numpy as np
import pandas as pd
import sys
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Define the path to the directory containing your modules
module_path = 'E:/GENERATING_DATA/ML_work'
if module_path not in sys.path:
    sys.path.append(module_path)

# Import the data retrieval function and configuration
from data_retrieval_R00 import main_data_retrieval
from config_R00 import GC_FILE

# Define the start and end dates for the data you want to retrieve
START_DATE = '2025-09-17 00:00'
END_DATE = '2025-09-21 23:00'
REPORT_OUTPUT_PATH = "actual_results.docx"

# Define the four heat duty columns
ENERGY_TAGS = [
    "HeatDuty_C-00_Reboiler",
    "HeatDuty_C-01_Reboiler",
    "HeatDuty_C-02_Reboiler",
    "HeatDuty_C-03_Reboiler"
]

def get_actual_energy_data(start_date, end_date):
    """
    Retrieves and sums the actual energy data for a given date range.
    """
    print(f"\n--- Retrieving Actual Data for {start_date} to {end_date} ---")
    
    scada_data, _ = main_data_retrieval(
        column_id=None,
        scada_tags=ENERGY_TAGS, 
        start_date=start_date, 
        end_date=end_date, 
        gc_file_path=GC_FILE, 
        gc_stream_ids=[]
    )

    if scada_data is None or scada_data.empty:
        print("🔴 ERROR: No data retrieved for the specified date range.")
        return None

    # Calculate the total energy by summing the four columns
    scada_data['TOTAL_ENERGY_KWH'] = scada_data[ENERGY_TAGS].sum(axis=1)

    # Clean and re-sample the data to a hourly frequency
    energy_df = scada_data[['DATEANDTIME', 'TOTAL_ENERGY_KWH']].copy()
    energy_df.set_index('DATEANDTIME', inplace=True)
    energy_df = energy_df.resample('H').mean().dropna()

    print("✅ Data retrieval and calculation complete.")
    return energy_df

def generate_actual_report(data_df, start_date, end_date, output_path):
    """Creates a detailed Word document report of the actual energy data."""
    doc = Document()
    doc.add_heading('Actual Energy Data Report', 0)
    
    doc.add_paragraph(f"Report Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Analysis Period: {start_date} to {end_date}")
    doc.add_paragraph("This report presents the actual historical energy data from the SCADA system for the specified period.")
    
    doc.add_heading('1. Hourly Energy Data', level=1)
    doc.add_paragraph("The table below shows the actual hourly energy values:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Actual Energy (kWh)'
    
    for date, energy in data_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(date.strftime('%Y-%m-%d %H:%M'))
        row_cells[1].text = f"{energy.values[0]:.2f}"
    
    doc.add_heading('2. Summary', level=1)
    total_energy_sum = data_df['TOTAL_ENERGY_KWH'].sum()
    doc.add_paragraph(f"The **total energy** for the period is: **{total_energy_sum:.2f} kWh**")
    
    doc.save(output_path)
    print(f"\n✅ Actual data report generated successfully at {output_path}")

if __name__ == "__main__":
    actual_data = get_actual_energy_data(START_DATE, END_DATE)
    
    if actual_data is not None:
        generate_actual_report(actual_data, START_DATE, END_DATE, REPORT_OUTPUT_PATH)
        print("\n--- Actual Energy Data (for comparison) ---")
        print(actual_data)